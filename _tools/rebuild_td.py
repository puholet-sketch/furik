# -*- coding: utf-8 -*-
"""Rebuild employment contract DOCX + PDF; copy into docs/трудовой-договор/.

Improvements vs prior draft:
- no disclaimer footer
- 5.3 bank details as label|value table (not one solid paragraph)
- section 10 starts on a new page; employer | worker in 2 columns
- black accents #1A1A1A, tighter spacing
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"D:\projects\Фурик")
sys.stdout.reconfigure(encoding="utf-8")

INK = RGBColor(0x1A, 0x1A, 0x1A)
TEXT = RGBColor(0x1A, 0x1A, 0x1A)

CONTRACT_STEM = "трудовой-договор-Ходжиматов-бессрочный-01.08.2026"
CONTRACT_DIR = ROOT / "трудовой-договор"
DOCS_DIR = ROOT / "docs" / "трудовой-договор"
DOCX_OUT = CONTRACT_DIR / f"{CONTRACT_STEM}.docx"
PDF_OUT = CONTRACT_DIR / f"{CONTRACT_STEM}.pdf"

WORKER = {
    "fio": "Ходжиматов Фуркатжон Махамаджонович",
    "fio_lat": "KHODJIMATOV FURKATJON MAKHAMADJONOVICH",
    "fio_short": "Ходжиматов Ф. М.",
    "dob": "01.02.1998",
    "passport": "FA6253664",
    "passport_issued": "29.08.2022, MIA 18227 (Узбекистан)",
    "passport_until": "28.08.2027",
    "inn": "772430430429",
    "snils": "229-173-224 63",
    "patent": "серия 77 № 2600347821",
    "patent_issued": "14.07.2026, Отдел внешней трудовой миграции УВМ ГУ МВД России по г. Москве",
    "patent_territory": "г. Москва",
    "patent_pr": "ПР 8114303",
    "patent_profession": "Бармен",
    "mig_series": "45 26",
    "mig_number": "0767385",
    "entry": "07.06.2026, КПП Внуково",
    "stay_until": "04.09.2026",
    "address": "г. Москва, поселение Вороновское, Варшавское ш. 64-й км, домовладение 1, строение 47",
    "phone": "8 (901) 797-57-53",
    "salary": "39 730",
    "salary_words": "тридцать девять тысяч семьсот тридцать",
}

EMPLOYER = {
    "fio": "Сорванова Анна Александровна",
    "fio_short": "Сорванова А. А.",
    "ogrnip": "322774600583080",
    "inn": "772973703990",
    "address": "г. Москва, ул. Василия Ланового, д. 3, кв. 243",
    "bank": 'МОСКОВСКИЙ ФИЛИАЛ АО КБ "МОДУЛЬБАНК"',
    "bik": "044525092",
    "ks": "30101810645250000092",
    "rs": "40802810370010393644",
    "email": "SORVANOVAAA@GMAIL.COM",
    "workplace": "г. Москва, ул. Киевская, д. 7, к. 2",
}

WORKER_BANK = {
    "currency": "RUB",
    "recipient": "ХОДЖИМАТОВ ФУРКАТЖОН МАХАМАДЖОНОВИЧ",
    "account": "40820810338110973759",
    "bank": "ПАО Сбербанк",
    "bik": "044525225",
    "ks": "30101810400000000225",
    "inn": "7707083893",
    "kpp": "773643001",
    "okpo": "57972160",
    "ogrn": "1027700132195",
    "swift": "SABRRUMM",
    "bank_address": "109544, МОСКВА, УЛ.Б.АНДРОНЬЕВСКАЯ,6",
}

CONTRACT_DATE = "01.08.2026"
START_DATE = "01.08.2026"


def set_run(run, *, size=11, bold=False, color=TEXT, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_bottom_border(paragraph, color_hex="1A1A1A", sz="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table, color="AAAAAA", sz="4") -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def set_nil_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def set_cell_margins(cell, top=40, bottom=40, left=40, right=40) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:tcMar"))
    if old is not None:
        tcPr.remove(old)
    tcMar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def clear_p(p) -> None:
    for child in list(p._p):
        if child.tag != qn("w:pPr"):
            p._p.remove(child)


def setup_doc(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(1.8)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.12
    pf.space_after = Pt(4)


def heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    set_run(r, size=12, bold=True, color=INK)
    add_bottom_border(p)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def para(
    doc: Document,
    text: str,
    *,
    bold=False,
    color=TEXT,
    size=11,
    align="justify",
    before=0,
    after=4,
) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_kv_table(
    doc_or_cell,
    rows: list[tuple[str, str]],
    *,
    label_width_cm=4.5,
    value_width_cm=11.0,
    font_size=10,
) -> None:
    if hasattr(doc_or_cell, "add_table"):
        table = doc_or_cell.add_table(rows=len(rows), cols=2)
    else:
        table = doc_or_cell.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, color="BBBBBB", sz="4")
    for i, (k, v) in enumerate(rows):
        cell0, cell1 = table.rows[i].cells
        cell0.width = Cm(label_width_cm)
        cell1.width = Cm(value_width_cm)
        set_cell_margins(cell0, 30, 30, 40, 40)
        set_cell_margins(cell1, 30, 30, 40, 40)
        clear_p(cell0.paragraphs[0])
        clear_p(cell1.paragraphs[0])
        p0, p1 = cell0.paragraphs[0], cell1.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r0 = p0.add_run(k)
        set_run(r0, size=font_size, bold=True, color=INK)
        r1 = p1.add_run(v)
        set_run(r1, size=font_size, color=TEXT)
        set_cell_shading(cell0, "F5F5F5")
        if i % 2 == 1:
            set_cell_shading(cell1, "FAFAFA")
    return table


def cell_para(cell, text: str, *, bold=False, size=10, center=False, before=0, after=4):
    if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text.strip():
        p = cell.paragraphs[0]
        clear_p(p)
    else:
        p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=INK)
    return p


def page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def build_contract(out: Path) -> Path:
    w, e, b = WORKER, EMPLOYER, WORKER_BANK
    doc = Document()
    setup_doc(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ТРУДОВОЙ ДОГОВОР")
    set_run(r, size=18, bold=True, color=INK)
    add_bottom_border(p, sz="18")
    p.paragraph_format.space_after = Pt(4)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(
        "с иностранным гражданином (Республика Узбекистан) · работа по патенту · на неопределённый срок"
    )
    set_run(r2, size=10, bold=True, color=INK)
    p2.paragraph_format.space_after = Pt(8)

    add_kv_table(
        doc,
        [
            ("Дата заключения", CONTRACT_DATE),
            ("Место заключения", "г. Москва"),
            ("Тип срока", "Бессрочный (на неопределённый срок)"),
            ("Дата начала работы", START_DATE),
            ("Должность", w["patent_profession"]),
            ("Оклад (до удержания НДФЛ)", f"{w['salary']} руб./мес."),
        ],
        label_width_cm=5.2,
        value_width_cm=11.0,
    )

    para(
        doc,
        f"Индивидуальный предприниматель {e['fio']}, ОГРНИП {e['ogrnip']}, "
        f"ИНН {e['inn']}, зарегистрированный по адресу: {e['address']}, "
        f"именуемый в дальнейшем «Работодатель», с одной стороны, и гражданин "
        f"Республики Узбекистан {w['fio']} ({w['fio_lat']}), дата рождения {w['dob']}, "
        f"паспорт № {w['passport']}, выдан {w['passport_issued']}, действует до {w['passport_until']}, "
        f"адрес пребывания: {w['address']}, ИНН {w['inn']}, СНИЛС {w['snils']}, "
        f"телефон {w['phone']}, именуемый в дальнейшем «Работник», с другой стороны, "
        f"вместе именуемые «Стороны», заключили настоящий трудовой договор "
        f"(далее — Договор) о нижеследующем.",
        before=8,
    )

    heading(doc, "1. Предмет договора")
    para(
        doc,
        f"1.1. Работодатель принимает Работника на работу на должность «{w['patent_profession']}», "
        f"а Работник обязуется лично выполнять трудовую функцию, соблюдать правила внутреннего "
        f"трудового распорядка и иные локальные акты Работодателя. "
        f"Должность соответствует профессии (виду деятельности), указанной в патенте Работника.",
    )
    para(
        doc,
        f"1.2. Место работы: {e['workplace']} (место осуществления деятельности Работодателя — кофейня).",
    )
    para(doc, "1.3. Работа по настоящему Договору является для Работника основной.")
    para(
        doc,
        f"1.4. Правовым основанием трудовой деятельности Работника на территории Российской Федерации "
        f"является патент ({w['patent']}), выданный {w['patent_issued']}. "
        f"Территория действия патента: {w['patent_territory']}. "
        f"Учётный номер (ПР): {w['patent_pr']}. "
        f"Профессия / вид деятельности по патенту: {w['patent_profession']}.",
    )
    para(
        doc,
        f"1.5. Миграционная карта: серия {w['mig_series']} № {w['mig_number']}, "
        f"въезд {w['entry']}, срок пребывания до {w['stay_until']}, цель визита — работа.",
    )

    heading(doc, "2. Срок договора")
    para(
        doc,
        f"2.1. Договор заключён на неопределённый срок. Дата начала работы: {START_DATE}. "
        f"Срок действия патента сам по себе не является основанием срочности трудового договора "
        f"(ч. 5 ст. 327.1 ТК РФ; письмо Роструда от 23.10.2013 № ПГ/9509-6-1).",
    )
    para(
        doc,
        "2.2. Работнику устанавливается испытательный срок продолжительностью 1 (один) месяц "
        "с даты начала работы. В период испытания на Работника распространяются положения "
        "трудового законодательства и иных актов, содержащих нормы трудового права.",
    )

    heading(doc, "3. Трудовые обязанности")
    para(doc, "3.1. В рамках должности «Бармен» Работник обязан:")
    duties = [
        "готовить кофейные, чайные и иные напитки на вынос по рецептурам и стандартам Работодателя;",
        "осуществлять продажу готовой продукции и сопутствующих товаров;",
        "принимать оплату от клиентов, соблюдать кассовую дисциплину и порядок учёта выручки "
        "по указанию Работодателя;",
        "консультировать клиентов по ассортименту напитков и продукции;",
        "поддерживать чистоту рабочей зоны, барной стойки, посуды и оборудования;",
        "соблюдать санитарные нормы и правила гигиены на рабочем месте;",
        "вести учёт расходных материалов, продуктов и инвентаря в объёме, "
        "определённом Работодателем;",
        "бережно относиться к имуществу Работодателя и сообщать о неисправностях оборудования;",
        "соблюдать требования охраны труда, пожарной безопасности и локальных актов Работодателя;",
        "своевременно сообщать об изменении паспортных, миграционных данных, реквизитов патента, "
        "СНИЛС, ИНН и банковских реквизитов;",
        "самостоятельно поддерживать действие патента (уплата фиксированных авансовых платежей "
        "по НДФЛ) и законность пребывания на территории РФ.",
    ]
    for i, d in enumerate(duties, start=1):
        para(doc, f"3.1.{i}. {d}")

    para(
        doc,
        "3.2. Работодатель обязан: предоставить работу, обусловленную Договором; "
        "обеспечить безопасные условия труда; своевременно и в полном объёме выплачивать "
        "заработную плату исключительно в безналичной форме; вести кадровый и налоговый учёт; "
        "уведомить территориальный орган МВД России о заключении и о расторжении настоящего "
        "Договора в срок не более 3 рабочих дней (п. 8 ст. 13 Федерального закона "
        "от 25.07.2002 № 115-ФЗ); соблюдать требования о персональных данных.",
    )
    para(
        doc,
        "3.3. Стороны имеют иные права и обязанности, предусмотренные ТК РФ, "
        "Законом № 115-ФЗ и иными нормативными правовыми актами РФ.",
    )

    heading(doc, "4. Режим рабочего времени и отдыха")
    para(
        doc,
        "4.1. Работнику устанавливается пятидневная рабочая неделя с двумя выходными днями — "
        "суббота и воскресенье. Нормальная продолжительность рабочего времени — "
        "40 (сорок) часов в неделю.",
    )
    para(
        doc,
        "4.2. Время присутствия на рабочем месте: с 08:00 до 18:00. "
        "В указанный интервал включается перерыв для отдыха и питания продолжительностью "
        "2 (два) часа — с 12:00 до 14:00, который не включается в рабочее время. "
        "Таким образом, продолжительность ежедневной работы составляет 8 часов "
        "(10 часов присутствия минус 2 часа перерыва), что соответствует норме "
        "40 часов в неделю.",
    )
    para(
        doc,
        "4.3. Ежегодный оплачиваемый отпуск предоставляется продолжительностью "
        "28 календарных дней в порядке, установленном ТК РФ.",
    )

    heading(doc, "5. Оплата труда")
    para(
        doc,
        f"5.1. Работнику устанавливается должностной оклад в размере {w['salary']} "
        f"({w['salary_words']}) рублей в месяц до удержания НДФЛ. "
        f"Сумма не ниже минимальной заработной платы в городе Москве на 2026 год "
        f"(39 730 руб. за полную норму рабочего времени).",
    )
    para(
        doc,
        "5.2. Заработная плата выплачивается исключительно в безналичном порядке путём "
        "перечисления на банковский счёт Работника. Выплата заработной платы наличными "
        "денежными средствами иностранному работнику — валютному нерезиденту не допускается "
        "(ст. 14 Федерального закона от 10.12.2003 № 173-ФЗ).",
    )
    para(doc, "5.3. Реквизиты для перечисления заработной платы:", align="left", after=2)
    add_kv_table(
        doc,
        [
            ("Получатель", b["recipient"]),
            ("Счёт", b["account"]),
            ("Банк", b["bank"]),
            ("БИК", b["bik"]),
            ("к/с", b["ks"]),
            ("Валюта", b["currency"]),
            ("ИНН банка", b["inn"]),
            ("КПП", b["kpp"]),
            ("ОКПО", b["okpo"]),
            ("ОГРН", b["ogrn"]),
            ("SWIFT", b["swift"]),
            ("Адрес банка", b["bank_address"]),
        ],
        label_width_cm=3.4,
        value_width_cm=12.8,
        font_size=9,
    )
    para(
        doc,
        "Выплата — только безналично на указанный счёт.",
        align="left",
        size=10,
        before=4,
    )
    para(
        doc,
        "5.4. Сроки выплаты заработной платы: аванс — до 15-го числа текущего месяца; "
        "окончательный расчёт — до 27-го числа текущего месяца. "
        "При совпадении дня выплаты с выходным или нерабочим праздничным днём выплата "
        "производится в предшествующий рабочий день (ст. 136 ТК РФ).",
    )
    para(
        doc,
        "5.5. Работодатель является налоговым агентом по НДФЛ. Работник самостоятельно "
        "уплачивает фиксированные авансовые платежи по патенту. Исчисленный НДФЛ может быть "
        "уменьшен на сумму таких авансов после получения уведомления налогового органа "
        "и заявления Работника с документами об уплате (ст. 227.1 НК РФ).",
    )
    para(
        doc,
        "5.6. При истечении, аннулировании патента либо неоплате фиксированного авансового "
        "платежа Работодатель отстраняет Работника от работы в порядке ст. 327.5 ТК РФ "
        "и при наличии оснований расторгает Договор.",
    )

    heading(doc, "6. Уведомление мвд и миграционные условия")
    para(
        doc,
        "6.1. Работодатель уведомляет территориальный орган МВД России о заключении "
        "и о расторжении Договора в срок не более 3 рабочих дней с даты соответствующего "
        "события (день события в срок не включается).",
    )
    para(
        doc,
        "6.2. Уведомление подаётся по форме, утверждённой приказом МВД России "
        "(до 01.09.2026 — приказ МВД от 30.07.2020 № 536 с изменениями; "
        "с 01.09.2026 — формы по приказу МВД от 12.05.2026 № 290), "
        "лично, почтой с описью вложения либо через портал Госуслуги.",
    )

    heading(doc, "7. Расторжение договора")
    para(
        doc,
        "7.1. Договор может быть прекращён по основаниям, предусмотренным ТК РФ, в том числе: "
        "по соглашению сторон (ст. 78); по инициативе Работника (ст. 80) — письменное заявление "
        "не позднее чем за две недели; по инициативе Работодателя (ст. 81) при наличии "
        "законных оснований.",
    )
    para(
        doc,
        "7.2. В день прекращения работы Работодатель выдаёт трудовую книжку / сведения "
        "о трудовой деятельности, производит окончательный расчёт и в срок не более "
        "3 рабочих дней уведомляет МВД о расторжении Договора.",
    )

    heading(doc, "8. Персональные данные")
    para(
        doc,
        "8.1. Работник даёт согласие на обработку персональных данных, необходимых "
        "для исполнения Договора, кадрового, налогового и миграционного учёта, "
        "в объёме и на срок, требуемые законом. Работодатель обеспечивает "
        "конфиденциальность и защиту персональных данных.",
    )

    heading(doc, "9. Заключительные положения")
    para(
        doc,
        "9.1. Договор составлен в двух экземплярах, имеющих одинаковую юридическую силу, "
        "по одному для каждой Стороны. Изменения оформляются дополнительными соглашениями "
        "в письменной форме.",
    )
    para(
        doc,
        f"9.2. Договор вступает в силу с {CONTRACT_DATE} и действует с даты начала работы "
        f"({START_DATE}).",
    )
    para(
        doc,
        "9.3. Во всём, что не урегулировано Договором, Стороны руководствуются ТК РФ, "
        "НК РФ, Федеральным законом № 115-ФЗ, Федеральным законом № 173-ФЗ "
        "и иными нормативными актами РФ.",
    )

    # --- Section 10: new page, 2 columns ---
    page_break(doc)
    heading(doc, "10. Реквизиты и подписи сторон")

    outer = doc.add_table(rows=1, cols=2)
    set_nil_borders(outer)
    left, right = outer.rows[0].cells
    left.width = Cm(8.6)
    right.width = Cm(8.6)
    set_cell_margins(left, 40, 40, 20, 80)
    set_cell_margins(right, 40, 40, 80, 20)

    cell_para(left, "РАБОТОДАТЕЛЬ", bold=True, size=11, center=True, after=6)
    add_kv_table(
        left,
        [
            ("Статус", f"ИП {e['fio']}"),
            ("ОГРНИП", e["ogrnip"]),
            ("ИНН", e["inn"]),
            ("Адрес", e["address"]),
            ("р/с", e["rs"]),
            ("Банк", e["bank"]),
            ("БИК", e["bik"]),
            ("к/с", e["ks"]),
            ("E-mail", e["email"]),
        ],
        label_width_cm=2.2,
        value_width_cm=5.8,
        font_size=8,
    )
    cell_para(left, f"_________________ / {e['fio_short']} /", size=8, before=10, after=2)
    cell_para(left, f"М.П. (при наличии)     Дата: {CONTRACT_DATE}", size=8, after=0)

    cell_para(right, "РАБОТНИК", bold=True, size=11, center=True, after=6)
    add_kv_table(
        right,
        [
            ("ФИО", w["fio"]),
            ("Гражданство", "Республика Узбекистан"),
            ("Паспорт", f"№ {w['passport']}, до {w['passport_until']}"),
            ("Адрес", w["address"]),
            ("Мигр. карта", f"серия {w['mig_series']} № {w['mig_number']}"),
            ("Патент", f"{w['patent']}, {w['patent_profession']}"),
            ("СНИЛС", w["snils"]),
            ("ИНН", w["inn"]),
            ("Телефон", w["phone"]),
            ("Счёт (безнал)", b["account"]),
            ("Банк", f"{b['bank']}, БИК {b['bik']}"),
        ],
        label_width_cm=2.6,
        value_width_cm=5.4,
        font_size=8,
    )
    cell_para(right, "Экземпляр Договора получил:", size=8, before=10, after=2)
    cell_para(right, f"_________________ / {w['fio_short']} /", size=8, after=2)
    cell_para(right, f"Дата: {CONTRACT_DATE}", size=8, after=0)

    # no disclaimer

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def export_pdf(docx_path: Path, pdf_path: Path) -> Path:
    import win32com.client  # type: ignore

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(docx_path.resolve()))
        # 17 = wdFormatPDF
        if pdf_path.exists():
            pdf_path.unlink()
        doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)
        doc.Close(False)
    finally:
        word.Quit()
    return pdf_path


def sync_docs(docx_path: Path, pdf_path: Path) -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(docx_path, DOCS_DIR / docx_path.name)
    shutil.copy2(pdf_path, DOCS_DIR / pdf_path.name)


def verify(docx_path: Path) -> None:
    doc = Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text

    asserts = {
        "no_disclaimer": "не является индивидуальной юридической консультацией" not in text,
        "no_disclaimer2": "Перед подписанием сверить" not in text,
        "has_barmen": "Бармен" in text,
        "has_date": "01.08.2026" in text,
        "has_kievskaya": "Киевская" in text,
        "has_account": "40820810338110973759" in text,
        "has_ip": "Сорванова" in text,
        "has_salary": "39 730" in text,
        "has_sec10": "10. РЕКВИЗИТЫ" in text.upper() or "10. Реквизиты" in text,
        "has_recipient_row": "ХОДЖИМАТОВ ФУРКАТЖОН" in text,
    }
    print("VERIFY", asserts)
    bad = [k for k, v in asserts.items() if not v]
    if bad:
        raise SystemExit(f"VERIFY_FAILED: {bad}")

    # page break before section 10
    found_break = False
    for p in doc.paragraphs:
        for r in p.runs:
            for br in r._element.findall(qn("w:br")):
                if br.get(qn("w:type")) == "page":
                    found_break = True
    print("HAS_PAGE_BREAK", found_break)
    if not found_break:
        raise SystemExit("missing page break before section 10")


def main() -> None:
    docx = build_contract(DOCX_OUT)
    print("WROTE", docx)
    verify(docx)
    pdf = export_pdf(docx, PDF_OUT)
    print("WROTE", pdf)
    sync_docs(docx, pdf)
    print("COPIED", DOCS_DIR / docx.name)
    print("COPIED", DOCS_DIR / pdf.name)


if __name__ == "__main__":
    main()
