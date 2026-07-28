# -*- coding: utf-8 -*-
"""Rebuild TD (black accents, compact requisites) + MVD hire notice DOCX."""
from __future__ import annotations

import hashlib
import json
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"D:\projects\Фурик")
sys.stdout.reconfigure(encoding="utf-8")

INK = RGBColor(0x1A, 0x1A, 0x1A)
TEXT = RGBColor(0x33, 0x33, 0x33)
MUTED = RGBColor(0x66, 0x66, 0x66)

CONTRACT_NAME = "трудовой-договор-Ходжиматов-бессрочный-01.08.2026.docx"
CONTRACT_OUT = ROOT / "трудовой-договор" / CONTRACT_NAME
NOTICE_DIR = ROOT / "уведомление-мвд"
NOTICE_OUT = NOTICE_DIR / "уведомление-МВД-заключение-ТД-Ходжиматов-01.08.2026.docx"

WORKER = {
    "fio": "Ходжиматов Фуркатжон Махамаджонович",
    "fio_lat": "KHODJIMATOV FURKATJON MAKHAMADJONOVICH",
    "fio_short": "Ходжиматов Ф. М.",
    "last": "Ходжиматов",
    "first": "Фуркатжон",
    "middle": "Махамаджонович",
    "dob": "01.02.1998",
    "passport": "FA6253664",
    "passport_issued": "29.08.2022, MIA 18227 (Узбекистан)",
    "passport_until": "28.08.2027",
    "inn": "772430430429",
    "snils": "229-173-224 63",
    "patent": "серия 77 № 2600347821",
    "patent_series": "77",
    "patent_number": "2600347821",
    "patent_issued": "14.07.2026, Отдел внешней трудовой миграции УВМ ГУ МВД России по г. Москве",
    "patent_issued_date": "14.07.2026",
    "patent_issuer": "Отдел внешней трудовой миграции УВМ ГУ МВД России по г. Москве",
    "patent_territory": "г. Москва",
    "patent_pr": "ПР 8114303",
    "patent_profession": "Бармен",
    "mig_series": "45 26",
    "mig_number": "0767385",
    "entry": "07.06.2026, КПП Внуково",
    "stay_until": "04.09.2026",
    "address": "г. Москва, поселение Вороновское, Варшавское ш. 64-й км, домовладение 1, строение 47",
    "phone": "8 (901) 797-57-53",
    "salary": "39 730",
    "salary_words": "тридцать девять тысяч семьсот тридцать",
}

EMPLOYER = {
    "fio": "Сорванова Анна Александровна",
    "fio_short": "Сорванова А. А.",
    "ogrnip": "322774600583080",
    "inn": "772973703990",
    "address": "г. Москва, ул. Василия Ланового, д. 3, кв. 243",
    "bank": 'МОСКОВСКИЙ ФИЛИАЛ АО КБ "МОДУЛЬБАНК"',
    "bik": "044525092",
    "ks": "30101810645250000092",
    "rs": "40802810370010393644",
    "email": "SORVANOVAAA@GMAIL.COM",
    "workplace": "г. Москва, ул. Киевская, д. 7, к. 2",
    # ОКВЭД уточнить по ЕГРИП; ориентир для кофейни
    "okved": "56.10",
}

WORKER_BANK = {
    "currency": "RUB",
    "recipient": "ХОДЖИМАТОВ ФУРКАТЖОН МАХАМАДЖОНОВИЧ",
    "account": "40820810338110973759",
    "bank": "ПАО Сбербанк",
    "bik": "044525225",
    "ks": "30101810400000000225",
    "inn": "7707083893",
    "kpp": "773643001",
    "okpo": "57972160",
    "ogrn": "1027700132195",
    "swift": "SABRRUMM",
    "bank_address": "109544, МОСКВА, УЛ.Б.АНДРОНЬЕВСКАЯ,6",
}

CONTRACT_DATE = "01.08.2026"
START_DATE = "01.08.2026"
CONTEXT_VERSION = 6


def set_run(run, *, size=11, bold=False, color=TEXT, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_bottom_border(paragraph, color_hex="1A1A1A", sz="18"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table, color="666666", sz="4") -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    # replace existing borders if any
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def setup_doc(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)


def heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    set_run(r, size=12, bold=True, color=INK)
    add_bottom_border(p, color_hex="1A1A1A", sz="12")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def para(
    doc: Document,
    text: str,
    *,
    bold=False,
    color=TEXT,
    size=11,
    align="justify",
) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(4)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_kv_table(doc: Document, rows: list[tuple[str, str]], *, label_width_cm=5.2) -> None:
    """Compact label|value table; left-aligned, no justify stretch."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, color="AAAAAA", sz="4")
    for i, (k, v) in enumerate(rows):
        cell0, cell1 = table.rows[i].cells
        cell0.text = ""
        cell1.text = ""
        p0 = cell0.paragraphs[0]
        p1 = cell1.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p0.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(k)
        set_run(r0, size=10, bold=True, color=INK)
        r1 = p1.add_run(v)
        set_run(r1, size=10, color=TEXT)
        set_cell_shading(cell0, "F5F5F5")
        if i % 2 == 1:
            set_cell_shading(cell1, "FAFAFA")
    # approximate column widths
    for row in table.rows:
        row.cells[0].width = Cm(label_width_cm)
        row.cells[1].width = Cm(11.5)


def worker_bank_block(compact: bool = False) -> str:
    b = WORKER_BANK
    if compact:
        return (
            f"валюта {b['currency']}; получатель — {b['recipient']}; "
            f"счёт {b['account']}; банк — {b['bank']}; БИК {b['bik']}; "
            f"к/с {b['ks']}; ИНН банка {b['inn']}; КПП {b['kpp']}; "
            f"ОКПО {b['okpo']}; ОГРН {b['ogrn']}; SWIFT {b['swift']}; "
            f"адрес банка — {b['bank_address']}"
        )
    return (
        f"валюта — {b['currency']}; "
        f"получатель — {b['recipient']}; "
        f"номер счёта — {b['account']}; "
        f"банк — {b['bank']}; "
        f"БИК — {b['bik']}; "
        f"корр. счёт — {b['ks']}; "
        f"ИНН банка — {b['inn']}; "
        f"КПП — {b['kpp']}; "
        f"ОКПО — {b['okpo']}; "
        f"ОГРН — {b['ogrn']}; "
        f"SWIFT — {b['swift']}; "
        f"почтовый адрес банка — {b['bank_address']}."
    )


def build_contract(out: Path) -> Path:
    w, e = WORKER, EMPLOYER
    doc = Document()
    setup_doc(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ТРУДОВОЙ ДОГОВОР")
    set_run(r, size=20, bold=True, color=INK)
    add_bottom_border(p, color_hex="1A1A1A")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(
        "с иностранным гражданином (Республика Узбекистан) · работа по патенту · на неопределённый срок"
    )
    set_run(r2, size=11, bold=True, color=INK)

    add_kv_table(
        doc,
        [
            ("Дата заключения", CONTRACT_DATE),
            ("Место заключения", "г. Москва"),
            ("Тип срока", "Бессрочный (на неопределённый срок)"),
            ("Дата начала работы", START_DATE),
            ("Должность", w["patent_profession"]),
            ("Оклад (до удержания НДФЛ)", f"{w['salary']} руб./мес."),
        ],
        label_width_cm=5.5,
    )

    para(
        doc,
        f"Индивидуальный предприниматель {e['fio']}, ОГРНИП {e['ogrnip']}, "
        f"ИНН {e['inn']}, зарегистрированный по адресу: {e['address']}, "
        f"именуемый в дальнейшем «Работодатель», с одной стороны, и гражданин "
        f"Республики Узбекистан {w['fio']} ({w['fio_lat']}), дата рождения {w['dob']}, "
        f"паспорт № {w['passport']}, выдан {w['passport_issued']}, действует до {w['passport_until']}, "
        f"адрес пребывания: {w['address']}, ИНН {w['inn']}, СНИЛС {w['snils']}, "
        f"телефон {w['phone']}, именуемый в дальнейшем «Работник», с другой стороны, "
        f"вместе именуемые «Стороны», заключили настоящий трудовой договор "
        f"(далее — Договор) о нижеследующем.",
    )

    heading(doc, "1. Предмет договора")
    para(
        doc,
        f"1.1. Работодатель принимает Работника на работу на должность «{w['patent_profession']}», "
        f"а Работник обязуется лично выполнять трудовую функцию, соблюдать правила внутреннего "
        f"трудового распорядка и иные локальные акты Работодателя. "
        f"Должность соответствует профессии (виду деятельности), указанной в патенте Работника.",
    )
    para(
        doc,
        f"1.2. Место работы: {e['workplace']} (место осуществления деятельности Работодателя — кофейня).",
    )
    para(doc, "1.3. Работа по настоящему Договору является для Работника основной.")
    para(
        doc,
        f"1.4. Правовым основанием трудовой деятельности Работника на территории Российской Федерации "
        f"является патент ({w['patent']}), выданный {w['patent_issued']}. "
        f"Территория действия патента: {w['patent_territory']}. "
        f"Учётный номер (ПР): {w['patent_pr']}. "
        f"Профессия / вид деятельности по патенту: {w['patent_profession']}.",
    )
    para(
        doc,
        f"1.5. Миграционная карта: серия {w['mig_series']} № {w['mig_number']}, "
        f"въезд {w['entry']}, срок пребывания до {w['stay_until']}, цель визита — работа.",
    )

    heading(doc, "2. Срок договора")
    para(
        doc,
        f"2.1. Договор заключён на неопределённый срок. Дата начала работы: {START_DATE}. "
        f"Срок действия патента сам по себе не является основанием срочности трудового договора "
        f"(ч. 5 ст. 327.1 ТК РФ; письмо Роструда от 23.10.2013 № ПГ/9509-6-1).",
    )
    para(
        doc,
        "2.2. Работнику устанавливается испытательный срок продолжительностью 1 (один) месяц "
        "с даты начала работы. В период испытания на Работника распространяются положения "
        "трудового законодательства и иных актов, содержащих нормы трудового права.",
    )

    heading(doc, "3. Трудовые обязанности")
    para(doc, "3.1. В рамках должности «Бармен» Работник обязан:")
    duties = [
        "готовить кофейные, чайные и иные напитки на вынос по рецептурам и стандартам Работодателя;",
        "осуществлять продажу готовой продукции и сопутствующих товаров;",
        "принимать оплату от клиентов, соблюдать кассовую дисциплину и порядок учёта выручки "
        "по указанию Работодателя;",
        "консультировать клиентов по ассортименту напитков и продукции;",
        "поддерживать чистоту рабочей зоны, барной стойки, посуды и оборудования;",
        "соблюдать санитарные нормы и правила гигиены на рабочем месте;",
        "вести учёт расходных материалов, продуктов и инвентаря в объёме, "
        "определённом Работодателем;",
        "бережно относиться к имуществу Работодателя и сообщать о неисправностях оборудования;",
        "соблюдать требования охраны труда, пожарной безопасности и локальных актов Работодателя;",
        "своевременно сообщать об изменении паспортных, миграционных данных, реквизитов патента, "
        "СНИЛС, ИНН и банковских реквизитов;",
        "самостоятельно поддерживать действие патента (уплата фиксированных авансовых платежей "
        "по НДФЛ) и законность пребывания на территории РФ.",
    ]
    for i, d in enumerate(duties, start=1):
        para(doc, f"3.1.{i}. {d}")

    para(
        doc,
        "3.2. Работодатель обязан: предоставить работу, обусловленную Договором; "
        "обеспечить безопасные условия труда; своевременно и в полном объёме выплачивать "
        "заработную плату исключительно в безналичной форме; вести кадровый и налоговый учёт; "
        "уведомить территориальный орган МВД России о заключении и о расторжении настоящего "
        "Договора в срок не более 3 рабочих дней (п. 8 ст. 13 Федерального закона "
        "от 25.07.2002 № 115-ФЗ); соблюдать требования о персональных данных.",
    )
    para(
        doc,
        "3.3. Стороны имеют иные права и обязанности, предусмотренные ТК РФ, "
        "Законом № 115-ФЗ и иными нормативными правовыми актами РФ.",
    )

    heading(doc, "4. Режим рабочего времени и отдыха")
    para(
        doc,
        "4.1. Работнику устанавливается пятидневная рабочая неделя с двумя выходными днями — "
        "суббота и воскресенье. Нормальная продолжительность рабочего времени — "
        "40 (сорок) часов в неделю.",
    )
    para(
        doc,
        "4.2. Время присутствия на рабочем месте: с 08:00 до 18:00. "
        "В указанный интервал включается перерыв для отдыха и питания продолжительностью "
        "2 (два) часа — с 12:00 до 14:00, который не включается в рабочее время. "
        "Таким образом, продолжительность ежедневной работы составляет 8 часов "
        "(10 часов присутствия минус 2 часа перерыва), что соответствует норме "
        "40 часов в неделю.",
    )
    para(
        doc,
        "4.3. Ежегодный оплачиваемый отпуск предоставляется продолжительностью "
        "28 календарных дней в порядке, установленном ТК РФ.",
    )

    heading(doc, "5. Оплата труда")
    para(
        doc,
        f"5.1. Работнику устанавливается должностной оклад в размере {w['salary']} "
        f"({w['salary_words']}) рублей в месяц до удержания НДФЛ. "
        f"Сумма не ниже минимальной заработной платы в городе Москве на 2026 год "
        f"(39 730 руб. за полную норму рабочего времени).",
    )
    para(
        doc,
        "5.2. Заработная плата выплачивается исключительно в безналичном порядке путём "
        "перечисления на банковский счёт Работника. Выплата заработной платы наличными "
        "денежными средствами иностранному работнику — валютному нерезиденту не допускается "
        "(ст. 14 Федерального закона от 10.12.2003 № 173-ФЗ).",
    )
    para(
        doc,
        "5.3. Реквизиты для перечисления заработной платы: "
        + worker_bank_block()
        + " Выплата — только безналично на указанный счёт.",
    )
    para(
        doc,
        "5.4. Сроки выплаты заработной платы: аванс — до 15-го числа текущего месяца; "
        "окончательный расчёт — до 27-го числа текущего месяца. "
        "При совпадении дня выплаты с выходным или нерабочим праздничным днём выплата "
        "производится в предшествующий рабочий день (ст. 136 ТК РФ).",
    )
    para(
        doc,
        "5.5. Работодатель является налоговым агентом по НДФЛ. Работник самостоятельно "
        "уплачивает фиксированные авансовые платежи по патенту. Исчисленный НДФЛ может быть "
        "уменьшен на сумму таких авансов после получения уведомления налогового органа "
        "и заявления Работника с документами об уплате (ст. 227.1 НК РФ).",
    )
    para(
        doc,
        "5.6. При истечении, аннулировании патента либо неоплате фиксированного авансового "
        "платежа Работодатель отстраняет Работника от работы в порядке ст. 327.5 ТК РФ "
        "и при наличии оснований расторгает Договор.",
    )

    heading(doc, "6. Уведомление мвд и миграционные условия")
    para(
        doc,
        "6.1. Работодатель уведомляет территориальный орган МВД России о заключении "
        "и о расторжении Договора в срок не более 3 рабочих дней с даты соответствующего "
        "события (день события в срок не включается).",
    )
    para(
        doc,
        "6.2. Уведомление подаётся по форме, утверждённой приказом МВД России "
        "(до 01.09.2026 — приказ МВД от 30.07.2020 № 536 с изменениями; "
        "с 01.09.2026 — формы по приказу МВД от 12.05.2026 № 290), "
        "лично, почтой с описью вложения либо через портал Госуслуги.",
    )

    heading(doc, "7. Расторжение договора")
    para(
        doc,
        "7.1. Договор может быть прекращён по основаниям, предусмотренным ТК РФ, в том числе: "
        "по соглашению сторон (ст. 78); по инициативе Работника (ст. 80) — письменное заявление "
        "не позднее чем за две недели; по инициативе Работодателя (ст. 81) при наличии "
        "законных оснований.",
    )
    para(
        doc,
        "7.2. В день прекращения работы Работодатель выдаёт трудовую книжку / сведения "
        "о трудовой деятельности, производит окончательный расчёт и в срок не более "
        "3 рабочих дней уведомляет МВД о расторжении Договора.",
    )

    heading(doc, "8. Персональные данные")
    para(
        doc,
        "8.1. Работник даёт согласие на обработку персональных данных, необходимых "
        "для исполнения Договора, кадрового, налогового и миграционного учёта, "
        "в объёме и на срок, требуемые законом. Работодатель обеспечивает "
        "конфиденциальность и защиту персональных данных.",
    )

    heading(doc, "9. Заключительные положения")
    para(
        doc,
        "9.1. Договор составлен в двух экземплярах, имеющих одинаковую юридическую силу, "
        "по одному для каждой Стороны. Изменения оформляются дополнительными соглашениями "
        "в письменной форме.",
    )
    para(
        doc,
        f"9.2. Договор вступает в силу с {CONTRACT_DATE} и действует с даты начала работы "
        f"({START_DATE}).",
    )
    para(
        doc,
        "9.3. Во всём, что не урегулировано Договором, Стороны руководствуются ТК РФ, "
        "НК РФ, Федеральным законом № 115-ФЗ, Федеральным законом № 173-ФЗ "
        "и иными нормативными актами РФ.",
    )

    # --- Реквизиты: компактные таблицы, LEFT, без justify ---
    heading(doc, "10. Реквизиты и подписи сторон")

    para(doc, "РАБОТОДАТЕЛЬ", bold=True, color=INK, size=11, align="left")
    add_kv_table(
        doc,
        [
            ("Статус", f"Индивидуальный предприниматель {e['fio']}"),
            ("ОГРНИП", e["ogrnip"]),
            ("ИНН", e["inn"]),
            ("Адрес", e["address"]),
            ("р/с", e["rs"]),
            ("Банк", e["bank"]),
            ("БИК", e["bik"]),
            ("к/с", e["ks"]),
            ("E-mail", e["email"]),
        ],
        label_width_cm=3.2,
    )
    para(
        doc,
        f"_________________ / {e['fio_short']} /     М.П. (при наличии)     Дата: {CONTRACT_DATE}",
        align="left",
        size=10,
    )

    para(doc, "РАБОТНИК", bold=True, color=INK, size=11, align="left")
    add_kv_table(
        doc,
        [
            ("ФИО", w["fio"]),
            ("Гражданство", "Республика Узбекистан"),
            ("Паспорт", f"№ {w['passport']}, выдан {w['passport_issued']}, до {w['passport_until']}"),
            ("Адрес пребывания", w["address"]),
            ("Миграционная карта", f"серия {w['mig_series']} № {w['mig_number']}"),
            (
                "Патент",
                f"{w['patent']}, профессия: {w['patent_profession']}, территория: {w['patent_territory']}",
            ),
            ("СНИЛС", w["snils"]),
            ("ИНН", w["inn"]),
            ("Телефон", w["phone"]),
            ("Зарплатный счёт (безнал)", worker_bank_block(compact=True)),
        ],
        label_width_cm=4.2,
    )
    para(
        doc,
        f"Экземпляр Договора получил: _________________ / {w['fio_short']} /     Дата: {CONTRACT_DATE}",
        align="left",
        size=10,
    )

    note = doc.add_paragraph()
    r = note.add_run(
        "Перед подписанием сверить данные с оригиналами документов и бухгалтером. "
        "Документ не является индивидуальной юридической консультацией."
    )
    set_run(r, size=9, color=MUTED)
    note.paragraph_format.space_before = Pt(12)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def build_mvd_notice(out: Path) -> Path:
    """Filled notice following App. 7 to MVD Order No. 536 (as amended)."""
    w, e = WORKER, EMPLOYER
    doc = Document()
    setup_doc(doc)
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("УВЕДОМЛЕНИЕ")
    set_run(r, size=14, bold=True, color=INK)

    para(
        doc,
        "о заключении трудового договора или гражданско-правового договора "
        "на выполнение работ (оказание услуг) с иностранным гражданином "
        "(лицом без гражданства)",
        bold=True,
        color=INK,
        size=11,
        align="center",
    )
    para(
        doc,
        "Форма: приложение № 7 к приказу МВД России от 30.07.2020 № 536 "
        "(в ред. приказов МВД от 22.11.2023 № 887 и от 06.08.2025 № 552). "
        "Действует до 01.09.2026; с этой даты — формы по приказу МВД от 12.05.2026 № 290.",
        size=9,
        color=MUTED,
        align="left",
    )

    para(doc, "Настоящее уведомление представляется в:", bold=True, align="left", size=10)
    add_kv_table(
        doc,
        [
            (
                "Территориальный орган МВД",
                "Управление по вопросам миграции ГУ МВД России по г. Москве "
                "(по месту фактической работы — г. Москва)",
            ),
        ],
        label_width_cm=5.5,
    )

    heading(doc, "1. Статус работодателя или заказчика работ (услуг)")
    para(doc, "[X] индивидуальный предприниматель", align="left", size=10)
    para(
        doc,
        "[ ] юридическое лицо   [ ] адвокат   [ ] физлицо — гражданин РФ   "
        "[ ] иное лицо   [ ] представительство / филиал иностранного юрлица   [ ] частный нотариус",
        align="left",
        size=9,
        color=MUTED,
    )
    add_kv_table(
        doc,
        [
            ("1.1. ОКВЭД (основной вид)", f"{e['okved']}  (сверить с ЕГРИП)"),
            ("1.2. Работодатель (ФИО ИП)", e["fio"]),
            ("ОГРНИП", e["ogrnip"]),
            ("ИНН", e["inn"]),
            ("Адрес места нахождения / регистрации", e["address"]),
            ("Контактный телефон", "________________ (указать телефон ИП)"),
            ("E-mail", e["email"]),
        ],
        label_width_cm=6.0,
    )

    heading(doc, "2. Сведения об иностранном гражданине")
    add_kv_table(
        doc,
        [
            ("2.1. Фамилия", w["last"]),
            ("2.2. Имя", w["first"]),
            ("2.3. Отчество", w["middle"]),
            ("2.4. Гражданство", "Республика Узбекистан"),
            ("2.6. Дата рождения", w["dob"]),
            ("2.7. Документ", "Паспорт иностранного гражданина"),
            ("Серия / номер", w["passport"]),
            ("Дата выдачи", "29.08.2022"),
            ("Кем выдан", "MIA 18227 (Узбекистан)"),
            ("Действует до", w["passport_until"]),
        ],
        label_width_cm=4.5,
    )

    heading(doc, "3. Сведения о разрешении на работу или патенте")
    add_kv_table(
        doc,
        [
            ("Наименование документа", "Патент"),
            ("Серия", w["patent_series"]),
            ("Номер", w["patent_number"]),
            ("Дата выдачи", w["patent_issued_date"]),
            ("Кем выдан", w["patent_issuer"]),
            (
                "Срок действия",
                f"с {w['patent_issued_date']} — по оплаченным периодам "
                f"(уточнить по патенту / квитанциям авансов; пребывание до {w['stay_until']})",
            ),
            ("Учётный номер (ПР)", w["patent_pr"]),
            ("Территория действия", w["patent_territory"]),
        ],
        label_width_cm=4.5,
    )
    para(
        doc,
        "3.1. Без разрешения/патента — не заполняется (деятельность на основании патента).",
        align="left",
        size=9,
        color=MUTED,
    )
    add_kv_table(
        doc,
        [
            ("3.2. Профессия / должность по договору", w["patent_profession"]),
            ("3.3. Основание", "[X] трудовой договор    [ ] гражданско-правовой договор"),
            ("Дата заключения договора", CONTRACT_DATE),
            ("3.4. Адрес места осуществления трудовой деятельности", e["workplace"]),
        ],
        label_width_cm=6.0,
    )

    heading(doc, "4. Подтверждение и подпись")
    para(
        doc,
        "Об ответственности за сообщение ложных сведений в уведомлении или предоставление "
        "поддельных документов предупреждён. С обработкой, передачей и хранением персональных "
        "данных согласен. Достоверность сведений, изложенных в настоящем уведомлении, подтверждаю:",
        align="left",
        size=10,
    )
    para(
        doc,
        f"Индивидуальный предприниматель _________________ / {e['fio_short']} /",
        align="left",
        size=10,
    )
    para(doc, f"Дата: «____» ______________ 20__ г.     М.П. (при наличии)", align="left", size=10)
    para(
        doc,
        "Уведомление подано по доверенности: № _____ от «____» ______________ 20__ г. "
        "(заполнить при подаче представителем)",
        align="left",
        size=9,
        color=MUTED,
    )

    note = doc.add_paragraph()
    r = note.add_run(
        "Это заполненный черновик по структуре официальной формы (прил. 7 к приказу МВД № 536). "
        "Перед подачей сверить ОКВЭД и срок патента с оригиналами; при личной/почтовой подаче "
        "рекомендуется перенести данные на актуальный бланк МВД (ячейки) либо подать через Госуслуги."
    )
    set_run(r, size=9, color=MUTED)
    note.paragraph_format.space_before = Pt(12)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def write_notice_readme(path: Path) -> None:
    path.write_text(
        """# Уведомление МВД о заключении ТД с иностранцем

## Срок
**Не позднее 3 рабочих дней** с даты заключения ТД (п. 8 ст. 13 Закона № 115-ФЗ).  
День заключения в срок **не входит**. Для ТД от **01.08.2026** край — **до конца 3-го рабочего дня после 01.08.2026**.

## Форма
- До **01.09.2026**: приложение № 7 к **приказу МВД от 30.07.2020 № 536** (ред. приказов № 887 и № 552).
- С **01.09.2026**: новые формы по **приказу МВД от 12.05.2026 № 290**.

Файл в этой папке — заполненный черновик по структуре прил. 7 (данные из ТД).

## Куда в Москве
Территориальный орган МВД **по месту фактической работы** (Киевская, 7к2 → **г. Москва**):  
**УВМ / ГУВМ ГУ МВД России по г. Москве** (актуальный адрес приёма — на сайте МВД / в окошке Госуслуг при подаче).

## Как подать
1. **Госуслуги** (удобнее): кабинет ИП + **УКЭП** → поиск «Уведомления по трудовым мигрантам» → уведомление о заключении ТД.  
   https://www.gosuslugi.ru/
2. **Лично** в подразделение по вопросам миграции (с отметкой о приёме).
3. **Почтой**: ценное письмо с описью вложения и уведомлением о вручении.

## Перед отправкой сверить
- ОКВЭД ИП по ЕГРИП (в черновике ориентир `56.10`).
- Срок действия патента по оригиналу / квитанциям авансов.
- Телефон ИП (в черновике оставлено поле).
""",
        encoding="utf-8",
    )


def file_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def verify_contract(path: Path) -> None:
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text

    colors = set()
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.color and r.font.color.rgb:
                colors.add(str(r.font.color.rgb))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if r.font.color and r.font.color.rgb:
                            colors.add(str(r.font.color.rgb))

    bad_colors = [c for c in colors if c.upper() in {"C00000", "FB5A5A"}]
    # requisites section: last tables should not be justify-only mega-gaps;
    # check that РАБОТОДАТЕЛЬ block uses tables
    assert len(doc.tables) >= 3, "expected meta + employer + worker tables"
    # alignment of requisites heading paragraphs
    req_paras = [p for p in doc.paragraphs if p.text.strip() in ("РАБОТОДАТЕЛЬ", "РАБОТНИК")]
    for p in req_paras:
        assert p.alignment in (WD_ALIGN_PARAGRAPH.LEFT, None) or p.alignment == WD_ALIGN_PARAGRAPH.LEFT

    checks = {
        "no_red": not bad_colors,
        "has_black_ink": "1A1A1A" in colors,
        "has_barmen": "Бармен" in text,
        "has_date": "01.08.2026" in text,
        "has_kievskaya": "Киевская" in text,
        "has_account": "40820810338110973759" in text,
        "has_ip": "Сорванова" in text and "322774600583080" in text,
        "tables": len(doc.tables),
        "colors": sorted(colors),
    }
    print("VERIFY_CONTRACT", json.dumps(checks, ensure_ascii=False))
    if bad_colors:
        raise SystemExit(f"Red colors remain: {bad_colors}")


def update_ai(contract_rel: str, notice_rel: str) -> None:
    context = ROOT / ".ai" / "CONTEXT.md"
    index = ROOT / ".ai" / "INDEX.md"
    state = ROOT / ".ai" / "state.json"
    today = "2026-07-28"

    context.write_text(
        f"""# Фурик — приём иностранца (Узбекистан)

**Цель:** оформить приём гражданина Узбекистана у ИП (~3 мес.): ТД, уведомления МВД, налоги/взносы.

**Статус:** бессрочный ТД пересобран (чёрные акценты, компактные реквизиты); черновик уведомления МВД о заключении ТД готов.

context_version: {CONTEXT_VERSION}  
updated: {today}

**Решения:**
- МЗП Москвы 2026 = **39 730 ₽**.
- Бессрочный ТД + увольнение по собственному (ст. 327.1 / 80 ТК РФ).
- Дата заключения и начала работы: **01.08.2026**.
- Должность в ТД: **Бармен** (как в патенте); обязанности — приготовление напитков/продажа/касса в рамках должности.
- Место работы (точка): **г. Москва, ул. Киевская, д. 7, к. 2**; юр.адрес ИП — ул. Василия Ланового (без изменений).
- График: пн–пт 08:00–18:00, перерыв 12:00–14:00 (2 ч), 40 ч/нед.; испытание 1 мес.
- ЗП: только безнал на счёт работника в ПАО Сбербанк; аванс до 15-го, расчёт до 27-го.
- Нагрузка ИП при окладе 39 730 ₽ (30% + травматизм 0,2%): **~51 728 ₽/мес**.
- Уведомление МВД: форма прил. 7 к приказу № 536 (до 01.09.2026); срок ≤3 раб. дня; Москва — по месту работы.

**Важно:** PII в `.ai/` не дублировать. Взносы/НДФЛ — сверить с бухгалтером. ОКВЭД в уведомлении — сверить с ЕГРИП.

**Следующий шаг:** подписать ТД; подать уведомление МВД ≤3 раб. дней (Госуслуги / лично / почта).
""",
        encoding="utf-8",
    )

    index.write_text(
        f"""# INDEX — Фурик

| Путь | Назначение |
|------|------------|
| `.ai/CONTEXT.md` | Краткий контекст |
| `.ai/state.json` | Машинное состояние |
| `.ai/summaries/legislation-foreign-worker-2026.md` | Справка по нормам 2026 |
| `документы-работника/` | Сжатые PDF: паспорт, патент, миграционная карта, квитанция аванса, ДМС, медосмотр |
| `{contract_rel}` | Актуальный бессрочный ТД (01.08.2026) |
| `{notice_rel}` | Черновик уведомления МВД о заключении ТД |
| `уведомление-мвд/README.md` | Срок 3 раб. дня, куда/как в Москве, Госуслуги |
| `инструкция/index.html` | HTML-инструкция в стиле Госуслуг + расчёт 2/3 мес. |
| `инструкция/styles.css` | Стили инструкции |
| `кадры-увольнение/заявление-об-увольнении.md` | Шаблон заявления об увольнении |
| `archive/` | Старые шаблоны ТД, md-инструкции, pdf_originals, compressed-копии, отчёты |
| `_tools/` | Скрипты генерации/сжатия/реорганизации |
""",
        encoding="utf-8",
    )

    active = {
        contract_rel: file_sha256(ROOT / contract_rel),
        notice_rel: file_sha256(ROOT / notice_rel),
        "уведомление-мвд/README.md": file_sha256(NOTICE_DIR / "README.md"),
        "инструкция/index.html": file_sha256(ROOT / "инструкция" / "index.html"),
        "инструкция/styles.css": file_sha256(ROOT / "инструкция" / "styles.css"),
        "кадры-увольнение/заявление-об-увольнении.md": file_sha256(
            ROOT / "кадры-увольнение" / "заявление-об-увольнении.md"
        ),
        ".ai/summaries/legislation-foreign-worker-2026.md": file_sha256(
            ROOT / ".ai" / "summaries" / "legislation-foreign-worker-2026.md"
        ),
    }
    state_obj = {
        "context_version": CONTEXT_VERSION,
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "context_hash": file_sha256(context),
        "active_files": list(active.keys()),
        "active_file_hashes": active,
        "decisions": [
            "Indefinite TD dated 01.08.2026 for Khodjimatov at IP Sorvanova AA",
            "Job title = patent profession Бармен; duties cover coffee-shop functions",
            "Workplace (point) = Moscow, Kievskaya 7 k.2; IP legal address remains Vasily Lanovoy",
            "Schedule Mon-Fri 08:00-18:00 with 2h break 12:00-14:00 (40h/week); probation 1 month",
            "Pay: non-cash only to Sber account 40820810338110973759; advance by 15th, final by 27th; salary 39730",
            "TD accents black (#1A1A1A); requisites as compact label|value tables",
            "MVD hire notice draft: App.7 Order 536; file within 3 business days via Gosuslugi/in person/mail",
        ],
        "blockers": [],
        "last_error": None,
        "retry_count": 0,
        "next_action": "Sign TD; submit MVD notice within 3 business days",
    }
    state.write_text(json.dumps(state_obj, ensure_ascii=False, indent=2), encoding="utf-8")
    print("AI_UPDATED", CONTEXT_VERSION)


def main() -> None:
    contract = build_contract(CONTRACT_OUT)
    print("WROTE", contract)
    verify_contract(contract)

    notice = build_mvd_notice(NOTICE_OUT)
    print("WROTE", notice)
    write_notice_readme(NOTICE_DIR / "README.md")
    print("WROTE", NOTICE_DIR / "README.md")

    contract_rel = f"трудовой-договор/{CONTRACT_NAME}"
    notice_rel = f"уведомление-мвд/{NOTICE_OUT.name}"
    update_ai(contract_rel, notice_rel)


if __name__ == "__main__":
    main()
