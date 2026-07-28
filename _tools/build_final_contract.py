# -*- coding: utf-8 -*-
"""Build final indefinite employment contract DOCX (black accents).

Prefer `_tools/rebuild_td_and_mvd.py` for TD + MVD notice rebuild.
"""
from __future__ import annotations

import hashlib
import json
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"D:\projects\Фурик")
sys.stdout.reconfigure(encoding="utf-8")

INK = RGBColor(0x1A, 0x1A, 0x1A)
CRIMSON = INK  # legacy alias — accents are black
EMBER = INK
TEXT = RGBColor(0x33, 0x33, 0x33)
MUTED = RGBColor(0x66, 0x66, 0x66)

CONTRACT_NAME = "трудовой-договор-Ходжиматов-бессрочный-01.08.2026.docx"
OUT = ROOT / "трудовой-договор" / CONTRACT_NAME

# Worker — from scans / previous filled TD / SNILS INN.txt
WORKER = {
    "fio": "Ходжиматов Фуркатжон Махамаджонович",
    "fio_lat": "KHODJIMATOV FURKATJON MAKHAMADJONOVICH",
    "fio_short": "Ходжиматов Ф. М.",
    "dob": "01.02.1998",
    "passport": "FA6253664",
    "passport_issued": "29.08.2022, MIA 18227 (Узбекистан)",
    "passport_until": "28.08.2027",
    "inn": "772430430429",
    "snils": "229-173-224 63",
    "patent": "серия 77 № 2600347821",
    "patent_issued": "14.07.2026, Отдел внешней трудовой миграции УВМ ГУ МВД России по г. Москве",
    "patent_territory": "г. Москва",
    "patent_pr": "ПР 8114303",
    "patent_profession": "Бармен",
    "mig_series": "45 26",
    "mig_number": "0767385",
    "entry": "07.06.2026, КПП Внуково",
    "stay_until": "04.09.2026",
    "address": "г. Москва, поселение Вороновское, Варшавское ш. 64-й км, домовладение 1, строение 47",
    "phone": "8 (901) 797-57-53",
    "salary": "39 730",
    "salary_words": "тридцать девять тысяч семьсот тридцать",
}

EMPLOYER = {
    "fio": "Сорванова Анна Александровна",
    "fio_short": "Сорванова А. А.",
    "ogrnip": "322774600583080",
    "inn": "772973703990",
    "address": "г. Москва, ул. Василия Ланового, д. 3, кв. 243",
    "bank": 'МОСКОВСКИЙ ФИЛИАЛ АО КБ "МОДУЛЬБАНК"',
    "bik": "044525092",
    "ks": "30101810645250000092",
    "rs": "40802810370010393644",
    "email": "SORVANOVAAA@GMAIL.COM",
    # Фактическое место работы (точка), не юр.адрес ИП
    "workplace": "г. Москва, ул. Киевская, д. 7, к. 2",
}

# Зарплатный счёт работника (RUB, только безнал)
WORKER_BANK = {
    "currency": "RUB",
    "recipient": "ХОДЖИМАТОВ ФУРКАТЖОН МАХАМАДЖОНОВИЧ",
    "account": "40820810338110973759",
    "bank": "ПАО Сбербанк",
    "bik": "044525225",
    "ks": "30101810400000000225",
    "inn": "7707083893",
    "kpp": "773643001",
    "okpo": "57972160",
    "ogrn": "1027700132195",
    "swift": "SABRRUMM",
    "bank_address": "109544, МОСКВА, УЛ.Б.АНДРОНЬЕВСКАЯ,6",
}

CONTRACT_DATE = "01.08.2026"
START_DATE = "01.08.2026"


def set_run(run, *, size=11, bold=False, color=TEXT, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_bottom_border(paragraph, color_hex="1A1A1A", sz="18"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def setup_doc(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)


def heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    set_run(r, size=12, bold=True, color=CRIMSON)
    add_bottom_border(p, sz="12")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


def para(doc: Document, text: str, *, bold=False, color=TEXT, size=11, align=None) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(4)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_meta_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        cell0, cell1 = table.rows[i].cells
        cell0.text = ""
        cell1.text = ""
        r0 = cell0.paragraphs[0].add_run(k)
        set_run(r0, size=10, bold=True, color=INK)
        r1 = cell1.paragraphs[0].add_run(v)
        set_run(r1, size=10, color=TEXT)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F5F5F5")
        shading.set(qn("w:val"), "clear")
        cell0._tc.get_or_add_tcPr().append(shading)
        if i % 2 == 1:
            shading2 = OxmlElement("w:shd")
            shading2.set(qn("w:fill"), "FAFAFA")
            shading2.set(qn("w:val"), "clear")
            cell1._tc.get_or_add_tcPr().append(shading2)


def worker_bank_block(compact: bool = False) -> str:
    b = WORKER_BANK
    if compact:
        return (
            f"валюта {b['currency']}; получатель — {b['recipient']}; "
            f"счёт {b['account']}; банк — {b['bank']}; БИК {b['bik']}; "
            f"к/с {b['ks']}; ИНН банка {b['inn']}; КПП {b['kpp']}; "
            f"ОКПО {b['okpo']}; ОГРН {b['ogrn']}; SWIFT {b['swift']}; "
            f"адрес банка — {b['bank_address']}"
        )
    return (
        f"валюта — {b['currency']}; "
        f"получатель — {b['recipient']}; "
        f"номер счёта — {b['account']}; "
        f"банк — {b['bank']}; "
        f"БИК — {b['bik']}; "
        f"корр. счёт — {b['ks']}; "
        f"ИНН банка — {b['inn']}; "
        f"КПП — {b['kpp']}; "
        f"ОКПО — {b['okpo']}; "
        f"ОГРН — {b['ogrn']}; "
        f"SWIFT — {b['swift']}; "
        f"почтовый адрес банка — {b['bank_address']}."
    )


def build_contract(out: Path) -> Path:
    w, e = WORKER, EMPLOYER
    doc = Document()
    setup_doc(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ТРУДОВОЙ ДОГОВОР")
    set_run(r, size=20, bold=True, color=INK)
    add_bottom_border(p)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(
        "с иностранным гражданином (Республика Узбекистан) · работа по патенту · на неопределённый срок"
    )
    set_run(r2, size=11, bold=True, color=EMBER)

    add_meta_table(
        doc,
        [
            ("Дата заключения", CONTRACT_DATE),
            ("Место заключения", "г. Москва"),
            ("Тип срока", "Бессрочный (на неопределённый срок)"),
            ("Дата начала работы", START_DATE),
            ("Должность", w["patent_profession"]),
            ("Оклад (до удержания НДФЛ)", f"{w['salary']} руб./мес."),
        ],
    )

    para(
        doc,
        f"Индивидуальный предприниматель {e['fio']}, ОГРНИП {e['ogrnip']}, "
        f"ИНН {e['inn']}, зарегистрированный по адресу: {e['address']}, "
        f"именуемый в дальнейшем «Работодатель», с одной стороны, и гражданин "
        f"Республики Узбекистан {w['fio']} ({w['fio_lat']}), дата рождения {w['dob']}, "
        f"паспорт № {w['passport']}, выдан {w['passport_issued']}, действует до {w['passport_until']}, "
        f"адрес пребывания: {w['address']}, ИНН {w['inn']}, СНИЛС {w['snils']}, "
        f"телефон {w['phone']}, именуемый в дальнейшем «Работник», с другой стороны, "
        f"вместе именуемые «Стороны», заключили настоящий трудовой договор "
        f"(далее — Договор) о нижеследующем.",
        align="justify",
    )

    heading(doc, "1. Предмет договора")
    para(
        doc,
        f"1.1. Работодатель принимает Работника на работу на должность «{w['patent_profession']}», "
        f"а Работник обязуется лично выполнять трудовую функцию, соблюдать правила внутреннего "
        f"трудового распорядка и иные локальные акты Работодателя. "
        f"Должность соответствует профессии (виду деятельности), указанной в патенте Работника.",
    )
    para(
        doc,
        f"1.2. Место работы: {e['workplace']} (место осуществления деятельности Работодателя — кофейня).",
    )
    para(
        doc,
        "1.3. Работа по настоящему Договору является для Работника основной.",
    )
    para(
        doc,
        f"1.4. Правовым основанием трудовой деятельности Работника на территории Российской Федерации "
        f"является патент ({w['patent']}), выданный {w['patent_issued']}. "
        f"Территория действия патента: {w['patent_territory']}. "
        f"Учётный номер (ПР): {w['patent_pr']}. "
        f"Профессия / вид деятельности по патенту: {w['patent_profession']}.",
    )
    para(
        doc,
        f"1.5. Миграционная карта: серия {w['mig_series']} № {w['mig_number']}, "
        f"въезд {w['entry']}, срок пребывания до {w['stay_until']}, цель визита — работа.",
    )

    heading(doc, "2. Срок договора")
    para(
        doc,
        f"2.1. Договор заключён на неопределённый срок. Дата начала работы: {START_DATE}. "
        f"Срок действия патента сам по себе не является основанием срочности трудового договора "
        f"(ч. 5 ст. 327.1 ТК РФ; письмо Роструда от 23.10.2013 № ПГ/9509-6-1).",
    )
    para(
        doc,
        "2.2. Работнику устанавливается испытательный срок продолжительностью 1 (один) месяц "
        "с даты начала работы. В период испытания на Работника распространяются положения "
        "трудового законодательства и иных актов, содержащих нормы трудового права.",
    )

    heading(doc, "3. Трудовые обязанности")
    para(
        doc,
        "3.1. В рамках должности «Бармен» Работник обязан:",
    )
    duties = [
        "готовить кофейные, чайные и иные напитки на вынос по рецептурам и стандартам Работодателя;",
        "осуществлять продажу готовой продукции и сопутствующих товаров;",
        "принимать оплату от клиентов, соблюдать кассовую дисциплину и порядок учёта выручки "
        "по указанию Работодателя;",
        "консультировать клиентов по ассортименту напитков и продукции;",
        "поддерживать чистоту рабочей зоны, барной стойки, посуды и оборудования;",
        "соблюдать санитарные нормы и правила гигиены на рабочем месте;",
        "вести учёт расходных материалов, продуктов и инвентаря в объёме, "
        "определённом Работодателем;",
        "бережно относиться к имуществу Работодателя и сообщать о неисправностях оборудования;",
        "соблюдать требования охраны труда, пожарной безопасности и локальных актов Работодателя;",
        "своевременно сообщать об изменении паспортных, миграционных данных, реквизитов патента, "
        "СНИЛС, ИНН и банковских реквизитов;",
        "самостоятельно поддерживать действие патента (уплата фиксированных авансовых платежей "
        "по НДФЛ) и законность пребывания на территории РФ.",
    ]
    for i, d in enumerate(duties, start=1):
        para(doc, f"3.1.{i}. {d}")

    para(
        doc,
        "3.2. Работодатель обязан: предоставить работу, обусловленную Договором; "
        "обеспечить безопасные условия труда; своевременно и в полном объёме выплачивать "
        "заработную плату исключительно в безналичной форме; вести кадровый и налоговый учёт; "
        "уведомить территориальный орган МВД России о заключении и о расторжении настоящего "
        "Договора в срок не более 3 рабочих дней (п. 8 ст. 13 Федерального закона "
        "от 25.07.2002 № 115-ФЗ); соблюдать требования о персональных данных.",
    )
    para(
        doc,
        "3.3. Стороны имеют иные права и обязанности, предусмотренные ТК РФ, "
        "Законом № 115-ФЗ и иными нормативными правовыми актами РФ.",
    )

    heading(doc, "4. Режим рабочего времени и отдыха")
    para(
        doc,
        "4.1. Работнику устанавливается пятидневная рабочая неделя с двумя выходными днями — "
        "суббота и воскресенье. Нормальная продолжительность рабочего времени — "
        "40 (сорок) часов в неделю.",
    )
    para(
        doc,
        "4.2. Время присутствия на рабочем месте: с 08:00 до 18:00. "
        "В указанный интервал включается перерыв для отдыха и питания продолжительностью "
        "2 (два) часа — с 12:00 до 14:00, который не включается в рабочее время. "
        "Таким образом, продолжительность ежедневной работы составляет 8 часов "
        "(10 часов присутствия минус 2 часа перерыва), что соответствует норме "
        "40 часов в неделю.",
    )
    para(
        doc,
        "4.3. Ежегодный оплачиваемый отпуск предоставляется продолжительностью "
        "28 календарных дней в порядке, установленном ТК РФ.",
    )

    heading(doc, "5. Оплата труда")
    para(
        doc,
        f"5.1. Работнику устанавливается должностной оклад в размере {w['salary']} "
        f"({w['salary_words']}) рублей в месяц до удержания НДФЛ. "
        f"Сумма не ниже минимальной заработной платы в городе Москве на 2026 год "
        f"(39 730 руб. за полную норму рабочего времени).",
    )
    para(
        doc,
        "5.2. Заработная плата выплачивается исключительно в безналичном порядке путём "
        "перечисления на банковский счёт Работника. Выплата заработной платы наличными "
        "денежными средствами иностранному работнику — валютному нерезиденту не допускается "
        "(ст. 14 Федерального закона от 10.12.2003 № 173-ФЗ).",
    )
    para(
        doc,
        "5.3. Реквизиты для перечисления заработной платы: "
        + worker_bank_block()
        + " Выплата — только безналично на указанный счёт.",
    )
    para(
        doc,
        "5.4. Сроки выплаты заработной платы: аванс — до 15-го числа текущего месяца; "
        "окончательный расчёт — до 27-го числа текущего месяца. "
        "При совпадении дня выплаты с выходным или нерабочим праздничным днём выплата "
        "производится в предшествующий рабочий день (ст. 136 ТК РФ).",
    )
    para(
        doc,
        "5.5. Работодатель является налоговым агентом по НДФЛ. Работник самостоятельно "
        "уплачивает фиксированные авансовые платежи по патенту. Исчисленный НДФЛ может быть "
        "уменьшен на сумму таких авансов после получения уведомления налогового органа "
        "и заявления Работника с документами об уплате (ст. 227.1 НК РФ).",
    )
    para(
        doc,
        "5.6. При истечении, аннулировании патента либо неоплате фиксированного авансового "
        "платежа Работодатель отстраняет Работника от работы в порядке ст. 327.5 ТК РФ "
        "и при наличии оснований расторгает Договор.",
    )

    heading(doc, "6. Уведомление мвд и миграционные условия")
    para(
        doc,
        "6.1. Работодатель уведомляет территориальный орган МВД России о заключении "
        "и о расторжении Договора в срок не более 3 рабочих дней с даты соответствующего "
        "события (день события в срок не включается).",
    )
    para(
        doc,
        "6.2. Уведомление подаётся по форме, утверждённой приказом МВД России "
        "(до 01.09.2026 — приказ МВД от 30.07.2020 № 536 с изменениями; "
        "с 01.09.2026 — формы по приказу МВД от 12.05.2026 № 290), "
        "лично, почтой с описью вложения либо через портал Госуслуги.",
    )

    heading(doc, "7. Расторжение договора")
    para(
        doc,
        "7.1. Договор может быть прекращён по основаниям, предусмотренным ТК РФ, в том числе: "
        "по соглашению сторон (ст. 78); по инициативе Работника (ст. 80) — письменное заявление "
        "не позднее чем за две недели; по инициативе Работодателя (ст. 81) при наличии "
        "законных оснований.",
    )
    para(
        doc,
        "7.2. В день прекращения работы Работодатель выдаёт трудовую книжку / сведения "
        "о трудовой деятельности, производит окончательный расчёт и в срок не более "
        "3 рабочих дней уведомляет МВД о расторжении Договора.",
    )

    heading(doc, "8. Персональные данные")
    para(
        doc,
        "8.1. Работник даёт согласие на обработку персональных данных, необходимых "
        "для исполнения Договора, кадрового, налогового и миграционного учёта, "
        "в объёме и на срок, требуемые законом. Работодатель обеспечивает "
        "конфиденциальность и защиту персональных данных.",
    )

    heading(doc, "9. Заключительные положения")
    para(
        doc,
        "9.1. Договор составлен в двух экземплярах, имеющих одинаковую юридическую силу, "
        "по одному для каждой Стороны. Изменения оформляются дополнительными соглашениями "
        "в письменной форме.",
    )
    para(
        doc,
        f"9.2. Договор вступает в силу с {CONTRACT_DATE} и действует с даты начала работы "
        f"({START_DATE}).",
    )
    para(
        doc,
        "9.3. Во всём, что не урегулировано Договором, Стороны руководствуются ТК РФ, "
        "НК РФ, Федеральным законом № 115-ФЗ, Федеральным законом № 173-ФЗ "
        "и иными нормативными актами РФ.",
    )

    heading(doc, "10. Реквизиты и подписи сторон")
    para(doc, "РАБОТОДАТЕЛЬ", bold=True, color=CRIMSON, size=11)
    para(
        doc,
        f"Индивидуальный предприниматель {e['fio']}\n"
        f"ОГРНИП: {e['ogrnip']}\n"
        f"ИНН: {e['inn']}\n"
        f"Адрес: {e['address']}\n"
        f"р/с {e['rs']}\n"
        f"Банк: {e['bank']}\n"
        f"БИК: {e['bik']}\n"
        f"к/с: {e['ks']}\n"
        f"E-mail: {e['email']}\n\n"
        f"_________________ / {e['fio_short']} /\n"
        f"М.П. (при наличии)\n"
        f"Дата: {CONTRACT_DATE}",
        size=10,
    )
    para(doc, "РАБОТНИК", bold=True, color=CRIMSON, size=11)
    para(
        doc,
        f"ФИО: {w['fio']}\n"
        f"Гражданство: Республика Узбекистан\n"
        f"Паспорт: № {w['passport']}, выдан {w['passport_issued']}, до {w['passport_until']}\n"
        f"Адрес пребывания: {w['address']}\n"
        f"Миграционная карта: серия {w['mig_series']} № {w['mig_number']}\n"
        f"Патент: {w['patent']}, профессия: {w['patent_profession']}, "
        f"территория: {w['patent_territory']}\n"
        f"СНИЛС: {w['snils']}\n"
        f"ИНН: {w['inn']}\n"
        f"Тел.: {w['phone']}\n"
        f"Реквизиты для перечисления заработной платы "
        f"(только безнал): {worker_bank_block(compact=True)}\n\n"
        f"Экземпляр Договора получил: _________________ / {w['fio_short']} /\n"
        f"Дата: {CONTRACT_DATE}",
        size=10,
    )

    note = doc.add_paragraph()
    r = note.add_run(
        "Перед подписанием сверить данные с оригиналами документов и бухгалтером. "
        "Документ не является индивидуальной юридической консультацией."
    )
    set_run(r, size=9, color=MUTED)
    note.paragraph_format.space_before = Pt(12)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def archive_old_contracts() -> list[str]:
    contract_dir = ROOT / "трудовой-договор"
    archive = ROOT / "archive"
    archive.mkdir(parents=True, exist_ok=True)
    moved = []
    for p in list(contract_dir.glob("*.docx")):
        if p.name == CONTRACT_NAME:
            continue
        dest = archive / p.name
        if dest.exists():
            stem, suf = dest.stem, dest.suffix
            n = 2
            while dest.exists():
                dest = archive / f"{stem}__old{n}{suf}"
                n += 1
        shutil.move(str(p), str(dest))
        moved.append(f"{p.name} → archive/{dest.name}")
    return moved


def update_instruction(new_rel: str) -> None:
    html = ROOT / "инструкция" / "index.html"
    if not html.exists():
        return
    text = html.read_text(encoding="utf-8")
    old = "трудовой-договор/трудовой-договор-Ходжиматов-бессрочный.docx"
    if old in text:
        text = text.replace(old, new_rel)
        html.write_text(text, encoding="utf-8")
        print("INSTR_UPDATED", new_rel)


def file_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def update_ai_context(contract_rel: str) -> None:
    context = ROOT / ".ai" / "CONTEXT.md"
    index = ROOT / ".ai" / "INDEX.md"
    state = ROOT / ".ai" / "state.json"

    context.write_text(
        """# Фурик — приём иностранца (Узбекистан)

**Цель:** оформить приём гражданина Узбекистана у ИП (~3 мес.): ТД, уведомления МВД, налоги/взносы.

**Статус:** бессрочный ТД готов к подписанию (01.08.2026); должность «Бармен»; место работы — Киевская; зарплатный счёт Сбер заполнен.

context_version: 5  
updated: 2026-07-28

**Решения:**
- МЗП Москвы 2026 = **39 730 ₽**.
- Бессрочный ТД + увольнение по собственному (ст. 327.1 / 80 ТК РФ).
- Дата заключения и начала работы: **01.08.2026**.
- Должность в ТД: **Бармен** (как в патенте); обязанности — приготовление напитков/продажа/касса в рамках должности.
- Место работы (точка): **г. Москва, ул. Киевская, д. 7, к. 2**; юр.адрес ИП — ул. Василия Ланового (без изменений).
- График: пн–пт 08:00–18:00, перерыв 12:00–14:00 (2 ч), 40 ч/нед.; испытание 1 мес.
- ЗП: только безнал на счёт работника в ПАО Сбербанк; аванс до 15-го, расчёт до 27-го.
- Нагрузка ИП при окладе 39 730 ₽ (30% + травматизм 0,2%): **~51 728 ₽/мес**.

**Важно:** PII в `.ai/` не дублировать. Взносы/НДФЛ — сверить с бухгалтером.

**Следующий шаг:** подписать ТД; уведомить МВД ≤3 раб. дней.
""",
        encoding="utf-8",
    )

    index.write_text(
        f"""# INDEX — Фурик

| Путь | Назначение |
|------|------------|
| `.ai/CONTEXT.md` | Краткий контекст |
| `.ai/state.json` | Машинное состояние |
| `.ai/summaries/legislation-foreign-worker-2026.md` | Справка по нормам 2026 |
| `документы-работника/` | Сжатые PDF: паспорт, патент, миграционная карта, квитанция аванса, ДМС, медосмотр |
| `{contract_rel}` | Актуальный бессрочный ТД (01.08.2026) |
| `инструкция/index.html` | HTML-инструкция в стиле Госуслуг + расчёт 2/3 мес. |
| `инструкция/styles.css` | Стили инструкции |
| `кадры-увольнение/заявление-об-увольнении.md` | Шаблон заявления об увольнении |
| `archive/` | Старые шаблоны ТД, md-инструкции, pdf_originals, compressed-копии, отчёты |
| `_tools/` | Скрипты генерации/сжатия/реорганизации |
""",
        encoding="utf-8",
    )

    context_hash = file_sha256(context)
    active = {
        contract_rel: file_sha256(ROOT / contract_rel),
        "инструкция/index.html": file_sha256(ROOT / "инструкция" / "index.html"),
        "инструкция/styles.css": file_sha256(ROOT / "инструкция" / "styles.css"),
        "кадры-увольнение/заявление-об-увольнении.md": file_sha256(
            ROOT / "кадры-увольнение" / "заявление-об-увольнении.md"
        ),
        ".ai/summaries/legislation-foreign-worker-2026.md": file_sha256(
            ROOT / ".ai" / "summaries" / "legislation-foreign-worker-2026.md"
        ),
    }
    state_obj = {
        "context_version": 5,
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "context_hash": context_hash,
        "active_files": list(active.keys()),
        "active_file_hashes": active,
        "decisions": [
            "Indefinite TD dated 01.08.2026 for Khodjimatov at IP Sorvanova AA",
            "Job title = patent profession Бармен; duties cover coffee-shop functions",
            "Workplace (point) = Moscow, Kievskaya 7 k.2; IP legal address remains Vasily Lanovoy",
            "Schedule Mon-Fri 08:00-18:00 with 2h break 12:00-14:00 (40h/week); probation 1 month",
            "Pay: non-cash only to Sber account 40820810338110973759; advance by 15th, final by 27th; salary 39730",
        ],
        "blockers": [],
        "last_error": None,
        "retry_count": 0,
        "next_action": "Sign TD; notify MVD within 3 business days",
    }
    state.write_text(json.dumps(state_obj, ensure_ascii=False, indent=2), encoding="utf-8")
    print("AI_UPDATED", context_hash[:16])


def verify_no_placeholders(path: Path) -> None:
    from docx import Document as D

    doc = D(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    bad = []
    for marker in ["заполнено по сканам", "В ПИСАТЬ", "[В ПИСАТЬ"]:
        if marker.lower() in text.lower() or marker in text:
            bad.append(marker)
    # Bank blank was exactly 8 underscores; signature lines keep longer runs
    if "Реквизиты для перечисления заработной платы: ________" in text:
        bad.append("salary bank still blank")
    if "Киевская" not in text:
        bad.append("missing Kievskaya workplace")
    if "Василия Ланового" not in text:
        bad.append("missing IP legal address")
    if "40820810338110973759" not in text:
        bad.append("missing worker account")
    if "бариста" in text.lower():
        bad.append("job title wrongly set to barista")
    print("VERIFY_MARKERS", bad if bad else "OK")
    print("HAS_BARMEN", "Бармен" in text)
    print("HAS_DATE", "01.08.2026" in text)
    print("HAS_SNILS", "229-173-224 63" in text)
    print("HAS_IP", "Сорванова" in text and "322774600583080" in text)
    print("HAS_WORKPLACE", "Киевская" in text)
    print("HAS_ACCOUNT", "40820810338110973759" in text)


def main() -> None:
    moved = archive_old_contracts()
    for m in moved:
        print("ARCHIVED", m)
    path = build_contract(OUT)
    print("WROTE", path)
    rel = f"трудовой-договор/{CONTRACT_NAME}"
    update_instruction(rel)
    update_ai_context(rel)
    verify_no_placeholders(path)
    # ensure only one docx in working folder
    remaining = list((ROOT / "трудовой-договор").glob("*.docx"))
    print("WORKING_CONTRACTS", [p.name for p in remaining])


if __name__ == "__main__":
    main()
