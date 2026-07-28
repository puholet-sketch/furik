# -*- coding: utf-8 -*-
"""Put employer/employee requisites side-by-side in section 10; tighten margins."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Length, Pt

ROOT = Path(r"D:\projects\Фурик")
DOCX = next(ROOT.glob("трудовой-договор/*.docx"))

sys.stdout.reconfigure(encoding="utf-8")


def set_cell_margins(cell, top=40, bottom=40, left=40, right=40) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:tcMar"))
    if old is not None:
        tcPr.remove(old)
    tcMar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_run_font(run, size_pt: float, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), "Times New Roman")


def clear_p(p) -> None:
    for child in list(p._p):
        if child.tag != qn("w:pPr"):
            p._p.remove(child)


def para(cell, text: str, *, bold=False, size=9, center=False, before=0, after=4):
    # Reuse empty first paragraph if present
    if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text.strip() and len(list(cell._tc.iterchildren())) == 1:
        p = cell.paragraphs[0]
        clear_p(p)
    else:
        p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size, bold=bold)
    return p


def fill_nested_table(parent_cell, rows_data: list[list[str]], label_cm: float, value_cm: float):
    tbl = parent_cell.add_table(rows=len(rows_data), cols=2)
    tbl.style = "Table Grid"
    label_w = Cm(label_cm)
    value_w = Cm(value_cm)
    for i, (label, value) in enumerate(rows_data):
        c0, c1 = tbl.rows[i].cells
        c0.width = label_w
        c1.width = value_w
        set_cell_margins(c0, 30, 30, 40, 40)
        set_cell_margins(c1, 30, 30, 40, 40)
        clear_p(c0.paragraphs[0])
        clear_p(c1.paragraphs[0])
        r0 = c0.paragraphs[0].add_run(label)
        set_run_font(r0, 8, bold=True)
        c0.paragraphs[0].paragraph_format.space_before = Pt(0)
        c0.paragraphs[0].paragraph_format.space_after = Pt(0)
        r1 = c1.paragraphs[0].add_run(value)
        set_run_font(r1, 8, bold=False)
        c1.paragraphs[0].paragraph_format.space_before = Pt(0)
        c1.paragraphs[0].paragraph_format.space_after = Pt(0)
    return tbl


def remove_element(el) -> None:
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def main() -> None:
    doc = Document(str(DOCX))

    sec = doc.sections[0]
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    employer_rows = [[c.text for c in row.cells] for row in doc.tables[1].rows]
    worker_rows = [[c.text for c in row.cells] for row in doc.tables[2].rows]

    body = doc.element.body
    children = list(body.iterchildren())

    idx_sec10 = None
    for i, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in child.iter(qn("w:t")))
        if text.startswith("10. РЕКВИЗИТЫ"):
            idx_sec10 = i
            break
    if idx_sec10 is None:
        raise SystemExit("section 10 not found")

    to_remove = []
    disclaimer_el = None
    for child in children[idx_sec10 + 1 :]:
        if child.tag == qn("w:sectPr"):
            break
        text = ""
        if child.tag == qn("w:p"):
            text = "".join(t.text or "" for t in child.iter(qn("w:t")))
        if text.startswith("Перед подписанием"):
            disclaimer_el = child
            break
        to_remove.append(child)

    for el in to_remove:
        remove_element(el)

    # Create outer table at end then move before disclaimer
    outer = doc.add_table(rows=1, cols=2)
    outer_tbl = outer._tbl
    body.remove(outer_tbl)
    if disclaimer_el is not None:
        disclaimer_el.addprevious(outer_tbl)
    else:
        sect_pr = body.find(qn("w:sectPr"))
        if sect_pr is not None:
            sect_pr.addprevious(outer_tbl)
        else:
            body.append(outer_tbl)

    # Switch to A4 portrait before sizing columns
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    usable = Length(int(sec.page_width) - int(sec.left_margin) - int(sec.right_margin))
    col_w = int(usable) // 2
    left, right = outer.rows[0].cells
    left.width = col_w
    right.width = col_w
    set_cell_margins(left, 40, 40, 40, 80)
    set_cell_margins(right, 40, 40, 80, 40)

    half_cm = usable.cm / 2 - 0.4
    label_l, value_l = 2.2, max(4.0, half_cm - 2.2)
    label_r, value_r = 2.6, max(4.0, half_cm - 2.6)

    para(left, "РАБОТОДАТЕЛЬ", bold=True, size=10, center=True, after=6)
    fill_nested_table(left, employer_rows, label_l, value_l)
    para(left, "_________________ / Сорванова А. А. /", size=8, before=8, after=2)
    para(left, "М.П. (при наличии)     Дата: 01.08.2026", size=8, before=0, after=0)

    para(right, "РАБОТНИК", bold=True, size=10, center=True, after=6)
    fill_nested_table(right, worker_rows, label_r, value_r)
    para(right, "Экземпляр Договора получил:", size=8, before=8, after=2)
    para(right, "_________________ / Ходжиматов Ф. М. /", size=8, before=0, after=2)
    para(right, "Дата: 01.08.2026", size=8, before=0, after=0)

    # Outer table: no borders, fixed width
    tblPr = outer_tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        outer_tbl.insert(0, tblPr)

    old_b = tblPr.find(qn("w:tblBorders"))
    if old_b is not None:
        tblPr.remove(old_b)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(int(usable)))
    tblW.set(qn("w:type"), "dxa")

    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    doc.save(str(DOCX))
    print("saved", DOCX.name)
    print(f"page A4 {sec.page_width.cm:.1f}x{sec.page_height.cm:.1f}, margins {sec.left_margin.cm} cm")
    # verify
    doc2 = Document(str(DOCX))
    print("tables:", len(doc2.tables))
    # last body order
    for child in list(doc2.element.body)[-8:]:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            t = "".join(x.text or "" for x in child.iter(qn("w:t")))[:70]
            print("P:", t)
        elif tag == "tbl":
            t = ""
            for x in child.iter(qn("w:t")):
                if x.text and x.text.strip():
                    t = x.text.strip()[:40]
                    break
            print("TBL:", t)


if __name__ == "__main__":
    main()
