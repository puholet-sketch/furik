# -*- coding: utf-8 -*-
"""Fill indefinite TD with known worker data and reorganize project folders."""
from __future__ import annotations

import hashlib
import json
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"D:\projects\Фурик")
sys.stdout.reconfigure(encoding="utf-8")

INK = RGBColor(0x1A, 0x1A, 0x1A)
CRIMSON = RGBColor(0xC0, 0x00, 0x00)
TEXT = RGBColor(0x33, 0x33, 0x33)
MUTED = RGBColor(0x66, 0x66, 0x66)

# Known from PDF scans (2026-07-28)
WORKER = {
    "fio": "Ходжиматов Фуркатжон Махамаджонович",
    "fio_lat": "KHODJIMATOV FURKATJON MAKHAMADJONOVICH",
    "dob": "01.02.1998",
    "passport": "FA6253664",
    "passport_issued": "29.08.2022, MIA 18227 (Узбекистан)",
    "passport_until": "28.08.2027",
    "inn": "772430430429",
    "patent": "серия 77 № 2600347821",
    "patent_issued": "14.07.2026, Отдел внешней трудовой миграции УВМ ГУ МВД России по г. Москве",
    "patent_territory": "г. Москва",
    "patent_pr": "ПР 8114303",
    "mig_series": "45 26",
    "mig_number": "0767385",
    "entry": "07.06.2026, КПП Внуково",
    "stay_until": "04.09.2026",
    "address": "г. Москва, поселение Вороновское, Варшавское ш. 64-й км, домовладение 1, строение 47",
    "phone": "8 (901) 797-57-53",
    "salary": "39 730",
    "salary_words": "тридцать девять тысяч семьсот тридцать",
}


def set_run(run, *, size=11, bold=False, color=TEXT, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_bottom_border(paragraph, color_hex="C00000", sz="18"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def setup_doc(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)


def heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    set_run(r, size=12, bold=True, color=CRIMSON)
    add_bottom_border(p, sz="12")
    p.paragraph_format.space_before = Pt(12)


def para(doc: Document, text: str, *, bold=False, color=TEXT, size=11) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(4)


def build_contract(out: Path) -> Path:
    w = WORKER
    doc = Document()
    setup_doc(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ТРУДОВОЙ ДОГОВОР")
    set_run(r, size=18, bold=True, color=INK)
    add_bottom_border(p)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(
        "с иностранным гражданином (на неопределённый срок — рекомендуемый вариант)"
    )
    set_run(r2, size=12, bold=True, color=CRIMSON)

    meta = doc.add_paragraph()
    r3 = meta.add_run(
        "г. Москва\t\t\t\t[В ПИСАТЬ: дата заключения ДД.ММ.ГГГГ]"
    )
    set_run(r3, size=11, color=MUTED)

    heading(doc, "1. Стороны договора")
    para(
        doc,
        "1.1. Работодатель: Индивидуальный предприниматель "
        "[В ПИСАТЬ: ФИО ИП полностью], ОГРНИП [В ПИСАТЬ: ОГРНИП], "
        "ИНН [В ПИСАТЬ: ИНН ИП], адрес регистрации: [В ПИСАТЬ: адрес ИП], "
        "именуемый в дальнейшем «Работодатель», с одной стороны.",
    )
    para(
        doc,
        f"1.2. Работник: гражданин Республики Узбекистан {w['fio']} "
        f"({w['fio_lat']}), дата рождения {w['dob']}, "
        f"паспорт: № {w['passport']}, выдан {w['passport_issued']}, "
        f"действует до {w['passport_until']}, "
        f"адрес пребывания: {w['address']}, "
        f"ИНН {w['inn']}, СНИЛС [В ПИСАТЬ: СНИЛС — в пакете документов отсутствует], "
        f"телефон {w['phone']}, "
        "именуемый в дальнейшем «Работник», с другой стороны.",
    )
    para(
        doc,
        "1.3. Работодатель и Работник совместно именуются «Стороны», а по отдельности — «Сторона».",
    )

    heading(doc, "2. Предмет договора")
    para(
        doc,
        "2.1. Работодатель принимает Работника на работу на должность "
        "«[В ПИСАТЬ: наименование должности — должно совпадать с профессией/видом "
        "деятельности в патенте]», а Работник обязуется лично выполнять трудовую "
        "функцию и соблюдать правила внутреннего трудового распорядка.",
    )
    para(
        doc,
        "2.2. Место работы: г. Москва, [В ПИСАТЬ: фактический адрес места работы].",
    )
    para(
        doc,
        "2.3. Договор заключён на неопределённый срок (ч. 5 ст. 327.1 ТК РФ). "
        "Срок действия патента сам по себе не является основанием срочности "
        "(письмо Роструда от 23.10.2013 № ПГ/9509-6-1).",
    )
    para(
        doc,
        "2.4. Дата начала работы: [В ПИСАТЬ: ДД.ММ.ГГГГ].",
    )
    para(
        doc,
        "2.5. Условие об испытании: [В ПИСАТЬ: без испытания / испытание "
        "продолжительностью … календарных дней].",
    )

    heading(doc, "3. Правовой статус иностранного работника")
    para(
        doc,
        "3.1. Работник осуществляет трудовую деятельность на основании патента "
        f"({w['patent']}), выданного {w['patent_issued']}. "
        f"Территория действия патента: {w['patent_territory']}. "
        f"Учётный номер (ПР): {w['patent_pr']}.",
    )
    para(
        doc,
        "3.2. Профессия / вид деятельности по патенту: "
        "[В ПИСАТЬ: как указано в патенте / приложении к патенту — сверить со сканом]. "
        "Фактическая работа должна соответствовать патенту.",
    )
    para(
        doc,
        f"3.3. Миграционная карта: серия {w['mig_series']} № {w['mig_number']}, "
        f"въезд {w['entry']}, срок пребывания до {w['stay_until']}, "
        "цель визита — работа (отмечено в карте).",
    )
    para(
        doc,
        "3.4. Работник самостоятельно уплачивает фиксированные авансовые платежи "
        "по НДФЛ за период действия патента (в Москве в 2026 ориентир ≈ 10 000 ₽/мес.; "
        "по квитанции от 19.07.2026 уже уплачено 10 000 ₽). "
        "Оплата патента — зона ответственности Работника.",
    )
    para(
        doc,
        "3.5. При истечении / аннулировании патента Работодатель обязан "
        "отстранить Работника от работы (ст. 327.5 ТК РФ) и при необходимости "
        "расторгнуть договор по установленным основаниям.",
    )

    heading(doc, "4. Права и обязанности")
    para(
        doc,
        "4.1. Работник обязан: добросовестно исполнять трудовую функцию; "
        "соблюдать ПВТР и требования охраны труда; своевременно сообщать "
        "об изменении паспортных / миграционных данных и статуса патента; "
        "поддерживать действие патента (оплата авансов).",
    )
    para(
        doc,
        "4.2. Работодатель обязан: предоставить работу, обусловленную договором; "
        "обеспечить безопасные условия труда; выплачивать заработную плату "
        "в установленные сроки только в безналичной форме; "
        "уведомить территориальный орган МВД о заключении и о расторжении "
        "настоящего договора в срок не позднее 3 рабочих дней "
        "(п. 8 ст. 13 Закона № 115-ФЗ).",
    )
    para(
        doc,
        "4.3. Должностные обязанности: "
        "[В ПИСАТЬ: кратко / со ссылкой на должностную инструкцию].",
    )

    heading(doc, "5. Оплата труда")
    para(
        doc,
        f"5.1. Работнику устанавливается должностной оклад в размере "
        f"{w['salary']} ({w['salary_words']}) рублей в месяц "
        "(не ниже минимальной заработной платы в г. Москве на 2026 год — 39 730 ₽).",
    )
    para(
        doc,
        "5.2. Заработная плата выплачивается исключительно в безналичном порядке "
        "на банковский счёт / карту Работника. Выплата наличными иностранному "
        "работнику не допускается.",
    )
    para(
        doc,
        "5.3. Банковские реквизиты для выплаты заработной платы: "
        "получатель — Ходжиматов Фуркатжон Махамаджонович; "
        "банк — [В ПИСАТЬ: наименование банка для зарплаты — в пакете есть квитанция "
        "АО КБ «Солидарность» об уплате аванса по патенту, но не реквизиты зарплатного счёта]; "
        "БИК [В ПИСАТЬ: БИК]; корр. счёт [В ПИСАТЬ: к/с]; "
        "расчётный счёт [В ПИСАТЬ: р/с]; "
        "номер карты (при необходимости) [В ПИСАТЬ: ****].",
    )
    para(
        doc,
        "5.4. Сроки выплаты: аванс — [В ПИСАТЬ: число] числа текущего месяца; "
        "окончательный расчёт — [В ПИСАТЬ: число] числа следующего месяца "
        "(не реже чем каждые полмесяца, ст. 136 ТК РФ).",
    )
    para(
        doc,
        "5.5. Работодатель является налоговым агентом по НДФЛ. "
        "Исчисленный НДФЛ может быть уменьшен на сумму фиксированных авансовых "
        "платежей по патенту после получения уведомления ИФНС и заявления Работника "
        "с квитанциями (ст. 227.1 НК РФ). Конкретные ставки и порядок — с бухгалтером.",
    )

    heading(doc, "6. Режим рабочего времени и отдыха")
    para(
        doc,
        "6.1. Работнику устанавливается [В ПИСАТЬ: пятидневная рабочая неделя / иной режим], "
        "нормальная продолжительность рабочего времени — 40 часов в неделю.",
    )
    para(
        doc,
        "6.2. Время начала работы: [В ПИСАТЬ: чч:мм]; окончания: [В ПИСАТЬ: чч:мм]; "
        "перерыв для отдыха и питания: [В ПИСАТЬ: с … по …].",
    )
    para(
        doc,
        "6.3. Ежегодный оплачиваемый отпуск — 28 календарных дней "
        "(ст. 115, 327.1 ТК РФ — на общих основаниях).",
    )

    heading(doc, "7. Изменение и прекращение договора")
    para(
        doc,
        "7.1. Изменение условий договора — по соглашению Сторон в письменной форме "
        "либо в иных случаях, предусмотренных ТК РФ.",
    )
    para(
        doc,
        "7.2. Договор может быть расторгнут по основаниям, предусмотренным ТК РФ, "
        "в том числе по инициативе Работника (ст. 80 ТК РФ — увольнение по собственному "
        "желанию с письменным предупреждением не менее чем за две недели).",
    )
    para(
        doc,
        "7.3. В день прекращения трудового договора Работодатель выдаёт Работнику "
        "трудовую книжку / сведения о трудовой деятельности и производит расчёт "
        "(ст. 84.1, 140 ТК РФ).",
    )
    para(
        doc,
        "7.4. О расторжении договора Работодатель уведомляет МВД в срок "
        "не позднее 3 рабочих дней.",
    )

    heading(doc, "8. Заключительные положения")
    para(
        doc,
        "8.1. Договор составлен в двух экземплярах, имеющих одинаковую юридическую силу, "
        "по одному для каждой из Сторон.",
    )
    para(
        doc,
        "8.2. Во всём, что не урегулировано настоящим Договором, Стороны "
        "руководствуются ТК РФ, Законом № 115-ФЗ и иными нормативными актами РФ.",
    )
    para(
        doc,
        "8.3. Договор вступает в силу с даты его подписания Сторонами "
        "(либо с даты начала работы — [В ПИСАТЬ: уточнить]).",
    )

    heading(doc, "9. Адреса, реквизиты и подписи сторон")
    para(
        doc,
        "РАБОТОДАТЕЛЬ:\n"
        "ИП [В ПИСАТЬ: ФИО]\n"
        "ОГРНИП [В ПИСАТЬ: …]  ИНН [В ПИСАТЬ: …]\n"
        "Адрес: [В ПИСАТЬ: …]\n"
        "р/с [В ПИСАТЬ: …] в [В ПИСАТЬ: банк], БИК [В ПИСАТЬ: …]\n"
        "Тел.: [В ПИСАТЬ: …]  E-mail: [В ПИСАТЬ: …]\n\n"
        "_________________ / [В ПИСАТЬ: ФИО ИП] /\n"
        "М.П. (при наличии)",
        bold=False,
    )
    para(
        doc,
        "РАБОТНИК:\n"
        f"{w['fio']}\n"
        f"Паспорт: {w['passport']}, выдан {w['passport_issued']}, до {w['passport_until']}\n"
        f"ИНН: {w['inn']}\n"
        f"Адрес: {w['address']}\n"
        f"Тел.: {w['phone']}\n"
        "Банковские реквизиты: [В ПИСАТЬ: как в п. 5.3]\n"
        "СНИЛС: [В ПИСАТЬ: СНИЛС]\n\n"
        f"Экземпляр трудового договора получил: _________________ / {w['fio']} /\n"
        "Дата: [В ПИСАТЬ: ДД.ММ.ГГГГ]",
    )

    note = doc.add_paragraph()
    r = note.add_run(
        "Заполнено по сканам из папки «документы-работника» на 2026-07-28. "
        "Поля «[В ПИСАТЬ: …]» — данные ИП / должность / зарплатный счёт / СНИЛС / даты. "
        "Перед подписанием сверить с оригиналами и бухгалтером. "
        "Не является индивидуальной юридической консультацией."
    )
    set_run(r, size=9, color=MUTED)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def move_safe(src: Path, dst: Path) -> None:
    if not src.exists():
        return
    dst.parent.mkdir(parents=True, exist_ok=True)
    if dst.exists():
        if src.is_file() and dst.is_file():
            if src.stat().st_size == dst.stat().st_size:
                src.unlink()
                return
            stem, suf = dst.stem, dst.suffix
            n = 2
            while dst.exists():
                dst = dst.with_name(f"{stem}__{n}{suf}")
                n += 1
        else:
            return
    shutil.move(str(src), str(dst))


def reorganize() -> None:
    docs_worker = ROOT / "документы-работника"
    contract_dir = ROOT / "трудовой-договор"
    instr_dir = ROOT / "инструкция"
    hr_dir = ROOT / "кадры-увольнение"
    archive = ROOT / "archive"
    tools_arch = archive / "tools-temp"

    for d in (docs_worker, contract_dir, instr_dir, hr_dir, archive, tools_arch):
        d.mkdir(parents=True, exist_ok=True)

    # Worker PDFs from root + compressed copies
    pdf_names = [
        "O'ZBEKISTON RESPUBLIKASI.pdf",
        "АО КБ Солидарность.pdf",
        "Б (ВыездDeparture).pdf",
        "Государственное_бюджетное_учреждение_здравоохранения_города.pdf",
        "МИНИСТЕРСТВО_ВНУТРЕННИХ_ДЕЛ_РОССИЙСКОЙ_ФЕДЕРАЦИИ.pdf",
        "СТРАХОВАЯ КОМПАНИЯ.pdf",
    ]
    # Move any PDF in root matching known names (encoding-safe via glob)
    for p in ROOT.glob("*.pdf"):
        move_safe(p, docs_worker / p.name)

    # compressed → archive (duplicates of working set)
    if (ROOT / "compressed").exists():
        move_safe(ROOT / "compressed", archive / "compressed")

    # originals → archive
    if (ROOT / "pdf_originals").exists():
        move_safe(ROOT / "pdf_originals", archive / "pdf_originals")

    # Old contracts / md → archive
    for name in [
        "Трудовой_договор_иностранец_Узбекистан.docx",
        "Трудовой_договор_черновик.md",
        "ИНСТРУКЦИЯ_прием_иностранца_Узбекистан.md",
        "PDF_COMPRESSION_REPORT.md",
        "PDF_COMPRESSION_REPORT.json",
        "PDF_сжатие_отчёт.md",
        "_file_inventory.json",
        "СНИЛС ИНН.txt",
    ]:
        move_safe(ROOT / name, archive / name)

    # docs/ folder contents
    docs = ROOT / "docs"
    if docs.exists():
        for p in docs.iterdir():
            if p.name.startswith("трудовой-договор"):
                move_safe(p, archive / "docs" / p.name)
            elif p.name == "заявление-об-увольнении.md":
                move_safe(p, hr_dir / p.name)
            elif p.name.startswith("инструкция"):
                move_safe(p, archive / "docs" / p.name)
            else:
                move_safe(p, archive / "docs" / p.name)
        try:
            docs.rmdir()
        except OSError:
            pass

    # OCR temp
    ocr = ROOT / "_tools" / "_ocr_preview"
    if ocr.exists():
        move_safe(ocr, tools_arch / "_ocr_preview")

    # Build filled contract
    contract_path = contract_dir / "трудовой-договор-Ходжиматов-бессрочный.docx"
    build_contract(contract_path)
    print("CONTRACT", contract_path)

    # Resignation template already moved; ensure exists
    resignation = hr_dir / "заявление-об-увольнении.md"
    if not resignation.exists():
        resignation.write_text(
            """# Заявление об увольнении по собственному желанию

*(распечатать / перенести на бланк; поля `[В ПИСАТЬ: …]` заполнить)*

---

Индивидуальному предпринимателю  
[В ПИСАТЬ: ФИО ИП]

от Ходжиматов Фуркатжон Махамаджонович  
должность: [В ПИСАТЬ: должность]

---

## Заявление

Прошу уволить меня по собственному желанию в соответствии со статьёй 80 Трудового кодекса Российской Федерации с **[В ПИСАТЬ: дата увольнения ДД.ММ.ГГГГ]**.

[Вариант при сокращении срока предупреждения по соглашению сторон:]  
Прошу расторгнуть трудовой договор **[В ПИСАТЬ: дата]** без отработки / с сокращением срока предупреждения по соглашению сторон.

Трудовой договор от **[В ПИСАТЬ: дата ТД]** № **[В ПИСАТЬ: номер, если есть]**.

Дата: **[В ПИСАТЬ: ДД.ММ.ГГГГ]**  

Подпись: _______________ / Ходжиматов Ф. М. /

---

### Для работодателя (памятка, не часть заявления)

- Исчисление 2 недель — с дня, следующего за днём получения заявления (ст. 80 ТК РФ).
- В день увольнения: приказ, расчёт, выдача документов.
- Уведомить МВД о расторжении ТД с иностранцем ≤ **3 рабочих дня**.
""",
            encoding="utf-8",
        )
    else:
        # Prefill worker name in existing template if still placeholder-only
        text = resignation.read_text(encoding="utf-8")
        if "Ходжиматов" not in text:
            text = text.replace(
                "от [В ПИСАТЬ: ФИО работника полностью]",
                "от Ходжиматов Фуркатжон Махамаджонович",
            )
            text = text.replace(
                "Подпись: _______________ / [В ПИСАТЬ: ФИО] /",
                "Подпись: _______________ / Ходжиматов Ф. М. /",
            )
            resignation.write_text(text, encoding="utf-8")
    print("RESIGNATION", resignation)


if __name__ == "__main__":
    reorganize()
    print("REORG_OK")
